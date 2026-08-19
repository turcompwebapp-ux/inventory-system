import streamlit as st
import pandas as pd
import gspread
from google.oauth2.service_account import Credentials
from datetime import datetime
import plotly.express as px
import plotly.graph_objects as go
import json
import uuid

st.set_page_config(
    page_title="TURCOMP Inventory",
    page_icon="📦",
    layout="wide",
    initial_sidebar_state="expanded"
)
# ── PASSWORD ──────────────────────────────────────────────────────────────────
def check_password():
    if "authenticated" not in st.session_state:
        st.session_state.authenticated = False

    if not st.session_state.authenticated:
        st.markdown("## 🔒 TURCOMP Inventory System")
        st.markdown("Please enter the access password to continue.")
        pwd = st.text_input("Password", type="password")
        if st.button("Login"):
            if pwd == st.secrets["APP_PASSWORD"]:
                st.session_state.authenticated = True
                st.rerun()
            else:
                st.error("❌ Incorrect password.")
        st.stop()

check_password()
# ── STYLING ───────────────────────────────────────────────────────────────────
st.markdown("""
<style>
[data-testid="stSidebar"] {
    background-color: #0f1e38;
}
[data-testid="stSidebar"] * {
    color: #cbd5e1 !important;
}
[data-testid="stSidebar"] .stRadio label {
    font-size: 15px !important;
    padding: 6px 0;
}
[data-testid="stSidebar"] [aria-checked="true"] + div {
    color: #ffffff !important;
    font-weight: 600 !important;
}
.kpi-card {
    background: #ffffff;
    border: 1px solid #e2e8f0;
    border-radius: 12px;
    padding: 18px 20px;
    display: flex;
    align-items: center;
    gap: 16px;
}
.kpi-icon {
    width: 48px; height: 48px;
    border-radius: 10px;
    display: flex; align-items: center; justify-content: center;
    font-size: 22px; flex-shrink: 0;
}
.kpi-label { font-size: 12px; color: #64748b; margin-bottom: 2px; }
.kpi-value { font-size: 26px; font-weight: 700; color: #0f1e38; line-height:1; }
.kpi-sub   { font-size: 11px; color: #94a3b8; margin-top: 2px; }
.cond-pill {
    border-radius: 10px; padding: 14px 10px;
    text-align: center; font-size: 13px; font-weight: 600;
}
.section-title {
    font-size: 15px; font-weight: 600;
    color: #0f1e38; margin: 18px 0 10px;
}
.chart-card {
    background: #ffffff; border: 1px solid #e2e8f0;
    border-radius: 12px; padding: 16px;
}
div[data-testid="stForm"] {
    border: 1px solid #e2e8f0;
    border-radius: 12px; padding: 20px;
    background: #ffffff;
}
.item-detail-card {
    background: #f8fafc; border: 1px solid #e2e8f0;
    border-radius: 10px; padding: 14px 18px; margin: 10px 0;
}
.stTabs [data-baseweb="tab"] {
    font-size: 14px; font-weight: 500;
}
</style>
""", unsafe_allow_html=True)

# ── CONSTANTS ─────────────────────────────────────────────────────────────────
SCOPES = [
    "https://www.googleapis.com/auth/spreadsheets",
    "https://www.googleapis.com/auth/drive"
]
CONDITION_OUT_OPTIONS   = ["GOOD", "DAMAGED (WORKING)"]
CONDITION_FULL_OPTIONS  = [
    "GOOD", "DAMAGED (WORKING)", "DAMAGED (NOT WORKING)",
    "UNDER REPAIR", "LOST", "DISPOSED", "SERVICE DUE"
]
RETURN_STATUS_OPTIONS = [
    "FULLY RETURNED", "PARTIALLY RETURNED", "LOST", "DAMAGED", "OUTSTANDING"
]
STATUS_OPTIONS = ["AVAILABLE", "OUT", "RETURNED", "PARTIALLY RETURNED", "UNDER MAINTENANCE", "LOST"]
COLOR_MAP_COND = {
    "GOOD":                  "#3B6D11",
    "DAMAGED":               "#E24B4A",
    "DAMAGED (NOT WORKING)": "#A32D2D",
    "DAMAGED (WORKING)":     "#BA7517",
    "LOST":                  "#534AB7",
    "MINOR ISSUE":           "#378ADD",
    "SERVICE DUE":           "#888780",
}
PILL_STYLE = {
    "GOOD":                  ("🟢", "#f0fdf4", "#166534"),
    "DAMAGED (WORKING)":     ("🟠", "#fff7ed", "#9a3412"),
    "DAMAGED (NOT WORKING)": ("🔴", "#fef2f2", "#991b1b"),
    "UNDER REPAIR":          ("🔵", "#eff6ff", "#1e40af"),
    "LOST":                  ("🟣", "#f5f3ff", "#5b21b6"),
    "DISPOSED":              ("⚫", "#f1f5f9", "#334155"),
    "SERVICE DUE":           ("⚪", "#f8fafc", "#475569"),
}

# ── GOOGLE SHEETS ─────────────────────────────────────────────────────────────
@st.cache_resource
def get_worksheet():
    creds = Credentials.from_service_account_info(
        st.secrets["gcp_service_account"],
        scopes=SCOPES
    )
    client = gspread.authorize(creds)
    return client.open("inventorydata").worksheet("inventorydata")
@st.cache_resource
def get_approvals_sheet():
    creds = Credentials.from_service_account_info(
        st.secrets["gcp_service_account"],
        scopes=SCOPES
    )
    client = gspread.authorize(creds)
    return client.open("inventorydata").worksheet("approvals")

@st.cache_data(ttl=15)
def load_approvals():
    ws   = get_approvals_sheet()
    data = ws.get_all_records()
    df = pd.DataFrame(data) if data else pd.DataFrame(columns=[
        "REQUEST_ID","KEYED_BY","DATE_REQUESTED","REQ_NAME","REQ_CONTACT",
        "REQ_POSITION","REQ_DEPT","REQ_PROJECT","REQ_JOBNO","REQ_LOCATION",
        "REQ_DATE","REQ_RETDATE","STATUS","APPROVED_BY","ITEMS_JSON","TYPE",
        "DELETED_BY","DELETE_REASON","RET_STATUS"
    ])
    if not df.empty and "STATUS" in df.columns:
        df = df[df["STATUS"] != "DELETED"]
    return df
def archive_old_approvals():
    """Move approved/rejected requests older than 30 days to an archive tab."""
    try:
        adf = load_approvals()
        if adf.empty:
            return
        aws = get_approvals_sheet()
        adf["DATE_REQUESTED"] = pd.to_datetime(
            adf["DATE_REQUESTED"], errors="coerce"
        )
        cutoff = pd.Timestamp.now() - pd.Timedelta(days=30)
        old = adf[
            (adf["STATUS"].isin(["APPROVED","REJECTED"])) &
            (adf["DATE_REQUESTED"] < cutoff)
        ]
        if old.empty:
            return

        # Try to get or create archive sheet
        try:
            archive_ws = get_worksheet().__class__
            creds = Credentials.from_service_account_info(
                st.secrets["gcp_service_account"], scopes=SCOPES
            )
            client = gspread.authorize(creds)
            spreadsheet = client.open("inventorydata")
            try:
                archive_ws = spreadsheet.worksheet("approvals_archive")
            except:
                archive_ws = spreadsheet.add_worksheet(
                    title="approvals_archive", rows=5000, cols=20
                )
                # Add headers
                aws_headers = aws.row_values(1)
                archive_ws.append_row(aws_headers)

            # Copy old rows to archive
            for _, row_data in old.iterrows():
                archive_ws.append_row(row_data.fillna("").astype(str).tolist())

            # Delete old rows from main approvals sheet (in reverse order)
            all_vals = aws.get_all_values()
            rows_to_delete = []
            for i, r in enumerate(all_vals[1:], start=2):
                if r and r[0] in old["REQUEST_ID"].values:
                    rows_to_delete.append(i)
            for row_idx in reversed(rows_to_delete):
                aws.delete_rows(row_idx)

            st.cache_data.clear()
        except Exception as e:
            pass  # Silently fail — don't block the app
    except:
        pass

def delete_approval_request(request_id, deleted_by, reason):
    """Marks a request as deleted — logged in sheet but hidden from app view."""
    aws = get_approvals_sheet()
    all_reqs = aws.get_all_values()
    for i, r in enumerate(all_reqs):
        if r and r[0] == request_id:
            aws.update(f"M{i+1}", [["DELETED"]])
            aws.update(f"Q{i+1}", [[deleted_by]])
            aws.update(f"R{i+1}", [[reason]])
            return True
    return False

@st.cache_resource
def get_settings_sheet():
    creds = Credentials.from_service_account_info(
        st.secrets["gcp_service_account"],
        scopes=SCOPES
    )
    client = gspread.authorize(creds)
    return client.open("inventorydata").worksheet("settings")

@st.cache_data(ttl=15)
def load_categories():
    ws = get_settings_sheet()
    values = ws.col_values(1)  # column A
    cats = [v.strip() for v in values[1:] if v.strip()]  # skip header row
    return sorted(cats)

def add_category(new_cat):
    ws = get_settings_sheet()
    existing = load_categories()
    if new_cat.strip().upper() in [c.upper() for c in existing]:
        return False
    ws.append_row([new_cat.strip().upper()])
    st.cache_data.clear()
    return True

def remove_category(cat_to_remove):
    ws = get_settings_sheet()
    all_values = ws.get_all_values()
    for i, row_data in enumerate(all_values):
        if row_data and row_data[0].strip().upper() == cat_to_remove.strip().upper():
            ws.delete_rows(i + 1)
            st.cache_data.clear()
            return True
    return False

@st.cache_data(ttl=30)
def load_data():
    ws   = get_worksheet()
    # Get raw string values to preserve leading zeros
    all_values = ws.get_all_values()
    if not all_values:
        return pd.DataFrame()
    headers = all_values[0]
    rows    = all_values[1:]
    df      = pd.DataFrame(rows, columns=headers)
    df      = df.replace("", None)
    df["QUANTITY"] = pd.to_numeric(df["QUANTITY"], errors="coerce").fillna(0)
    return df

def reload():
    st.cache_data.clear()
    st.rerun()

# ── SEARCH HELPER ─────────────────────────────────────────────────────────────
def search_item(df, label="search"):
    search_by = st.selectbox(
        "Search by",
        ["Serial Number", "Tagging Number", "Description", "Brand", "Category"],
        key=f"sb_{label}"
    )
    query = st.text_input(
        f"Enter {search_by} (or part of it)",
        key=f"q_{label}"
    )
    if not query:
        return None, None

    col_map = {
        "Serial Number":  "SERIAL NUMBER",
        "Tagging Number": "TAGGING NUMBER",
        "Description":    "DESCRIPTION",
        "Brand":          "BRAND",
        "Category":       "CATEGORY",
    }
    col     = col_map[search_by]
    matches = df[
        df[col].fillna("").astype(str).str.strip().str.upper()
        .str.contains(query.strip().upper(), regex=False)
    ]

    if matches.empty:
        st.error(f"❌ No items found matching '{query}' in {search_by}.")
        return None, None

    # Optional category filter if many results
    if len(matches) > 5:
        cats = ["All"] + sorted(matches["CATEGORY"].dropna().unique().tolist())
        cat_filter = st.selectbox("Narrow by Category", cats, key=f"cat_{label}")
        if cat_filter != "All":
            matches = matches[matches["CATEGORY"] == cat_filter]

    if len(matches) == 0:
        st.error("No items found after filtering.")
        return None, None

    # Build readable option labels
    def make_label(r):
        parts = []
        for field in ["SERIAL NUMBER", "TAGGING NUMBER", "DESCRIPTION", "BRAND", "CATEGORY", "SIZE"]:
            val = r.get(field)
            if val and str(val).strip() not in ["-", "None", ""]:
                parts.append(f"{field.title()}: {val}")
        return " | ".join(parts) if parts else f"Row {r.name + 2}"

    options    = matches.apply(make_label, axis=1).tolist()
    chosen     = st.selectbox(
        f"{len(matches)} item(s) found — select one",
        options, key=f"pick_{label}"
    )
    idx        = options.index(chosen)
    row        = matches.iloc[idx]
    row_num    = int(matches.index[idx]) + 2

    # Detail card
    st.markdown(f"""
    <div class="item-detail-card">
        <div style="display:grid;grid-template-columns:1fr 1fr;gap:6px;font-size:13px;">
            <div><b>Description:</b> {row.get('DESCRIPTION') or '—'}</div>
            <div><b>Brand:</b> {row.get('BRAND') or '—'}</div>
            <div><b>Type/Spec:</b> {row.get('TYPE/SPEC') or '—'}</div>
            <div><b>Size:</b> {row.get('SIZE') or '—'}</div>
            <div><b>Serial No:</b> {row.get('SERIAL NUMBER') or '—'}</div>
            <div><b>Tagging No:</b> {row.get('TAGGING NUMBER') or '—'}</div>
            <div><b>Category:</b> {row.get('CATEGORY') or '—'}</div>
            <div><b>Remarks:</b> {row.get('REMARKS') or '—'}</div>
        </div>
    </div>
    """, unsafe_allow_html=True)

    return row, row_num
# ── PDF GENERATORS ────────────────────────────────────────────────────────────
def generate_issue_pdf(data):
    from reportlab.lib.pagesizes import A4
    from reportlab.lib import colors
    from reportlab.lib.units import mm
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.enums import TA_CENTER, TA_LEFT
    import io

    buf    = io.BytesIO()
    doc    = SimpleDocTemplate(buf, pagesize=A4,
                               topMargin=10*mm, bottomMargin=10*mm,
                               leftMargin=15*mm, rightMargin=15*mm)
    styles = getSampleStyleSheet()
    story  = []

    bold   = ParagraphStyle("bold", parent=styles["Normal"], fontName="Helvetica-Bold")
    center = ParagraphStyle("center", parent=styles["Normal"], alignment=TA_CENTER)
    title  = ParagraphStyle("title", parent=styles["Normal"],
                             fontName="Helvetica-Bold", fontSize=11, alignment=TA_CENTER)
    small  = ParagraphStyle("small", parent=styles["Normal"], fontSize=8)

    # Header table with logo
    import os
    from reportlab.platypus import Image as RLImage

    logo_path = "logo.png"
    if os.path.exists(logo_path):
        logo_cell = RLImage(logo_path, width=20*mm, height=20*mm)
    else:
        logo_cell = Paragraph("", small)

    header_data = [[
        logo_cell,
        Paragraph("<b>TURCOMP ENGINEERING SERVICES SDN. BHD.</b>", title),
        Paragraph("Form No.", small),
        Paragraph("Rev. No.", small),
    ],[
        Paragraph("", small),
        Paragraph("<b>WAREHOUSE ISSUE FORM</b>", title),
        Paragraph("Issue Date:", small),
        Paragraph("Page No.", small),
    ]]
    ht = Table(header_data, colWidths=[22*mm, 98*mm, 30*mm, 30*mm])
    ht.setStyle(TableStyle([
        ("BOX",        (0,0), (-1,-1), 0.5, colors.black),
        ("INNERGRID",  (0,0), (-1,-1), 0.5, colors.black),
        ("VALIGN",     (0,0), (-1,-1), "MIDDLE"),
        ("ALIGN",      (0,0), (0,-1), "CENTER"),
        ("SPAN",       (0,0), (0,1)),
        ("ROWBACKGROUNDS", (0,0), (-1,-1), [colors.white]),
    ]))
    story.append(ht)
    story.append(Spacer(1, 4*mm))

    # Section A — Requestor Info
    story.append(Paragraph("<b>A. REQUESTOR INFORMATION</b>", bold))
    story.append(Spacer(1, 2*mm))
    req_data = [
        ["Name", data.get("req_name",""), "Contact No", data.get("req_contact","")],
        ["Position", data.get("req_position",""), "Department", data.get("req_dept","")],
        ["Project Name", data.get("req_project",""), "Project / Job No", data.get("req_jobno","")],
        ["Usage Location", data.get("req_location",""), "", ""],
        ["Date Requested", data.get("req_date",""), "Required Return Date", data.get("req_retdate","")],
    ]
    rt = Table(req_data, colWidths=[35*mm, 65*mm, 40*mm, 40*mm])
    rt.setStyle(TableStyle([
        ("BOX",       (0,0), (-1,-1), 0.5, colors.black),
        ("INNERGRID", (0,0), (-1,-1), 0.5, colors.black),
        ("FONTNAME",  (0,0), (0,-1), "Helvetica-Bold"),
        ("FONTNAME",  (2,0), (2,-1), "Helvetica-Bold"),
        ("FONTSIZE",  (0,0), (-1,-1), 8),
        ("VALIGN",    (0,0), (-1,-1), "MIDDLE"),
        ("SPAN",      (1,3), (3,3)),
    ]))
    story.append(rt)
    story.append(Spacer(1, 4*mm))

    # Section B — Asset Details
    story.append(Paragraph("<b>B. ASSET ISSUE DETAILS</b>", bold))
    story.append(Spacer(1, 2*mm))

    asset_header = [["No", "Asset Tag No.", "Description", "Serial No.",
                      "Qty Issued", "Date Out", "Condition Out", "Status"]]
    items = data.get("items", [])
    for i, item in enumerate(items, 1):
        asset_header.append([
            str(i),
            str(item.get("TAGGING NUMBER","") or "-"),
            str(item.get("DESCRIPTION","") or "-"),
            str(item.get("SERIAL NUMBER","") or "-"),
            "1",
            data.get("date_out",""),
            data.get("cond_out","GOOD"),
            "Issued"
        ])
    # Add blank rows up to 10
    while len(asset_header) < 11:
        asset_header.append(["", "", "", "", "", "", "", ""])

    at = Table(asset_header, colWidths=[10*mm, 30*mm, 45*mm, 30*mm, 15*mm, 22*mm, 18*mm, 10*mm])
    at.setStyle(TableStyle([
        ("BOX",         (0,0), (-1,-1), 0.5, colors.black),
        ("INNERGRID",   (0,0), (-1,-1), 0.5, colors.black),
        ("BACKGROUND",  (0,0), (-1,0),  colors.HexColor("#1e3a5f")),
        ("TEXTCOLOR",   (0,0), (-1,0),  colors.white),
        ("FONTNAME",    (0,0), (-1,0),  "Helvetica-Bold"),
        ("FONTSIZE",    (0,0), (-1,-1), 7.5),
        ("ALIGN",       (0,0), (-1,-1), "CENTER"),
        ("VALIGN",      (0,0), (-1,-1), "MIDDLE"),
        ("ROWHEIGHT",   (0,1), (-1,-1), 8*mm),
    ]))
    story.append(at)
    story.append(Spacer(1, 6*mm))

    # Section E — Issue Sign-Off
    story.append(Paragraph("<b>E. ISSUE SIGN-OFF</b>", bold))
    story.append(Spacer(1, 2*mm))
    keyed_by    = data.get("keyed_by", "")
    approved_by = data.get("approved_by", "")
    now_str     = datetime.now().strftime("%Y-%m-%d %H:%M")
    sign_data = [
        [Paragraph("<b>Prepared by<br/>(Storekeeper)</b>", small),
         Paragraph("<b>Checked &amp; approved by<br/>(Material Controller/Warehouse Coordinator)</b>", small),
         Paragraph("<b>Received by (Requestor/Receiver)</b>", small)],
        [Paragraph(f"Electronically signed by: <b>{keyed_by}</b><br/>Date: {now_str}", small),
         Paragraph(f"Electronically approved by: <b>{approved_by}</b><br/>Date: {now_str}", small),
         Paragraph(
             "Name: ________________"
             "<br/><br/><br/><br/>"
             "Signature: ________________"
             "<br/><br/><br/>"
             "Date: ___________",
             small
         )],
    ]
    st_table = Table(sign_data, colWidths=[60*mm, 60*mm, 60*mm])
    st_table.setStyle(TableStyle([
        ("BOX",       (0,0), (-1,-1), 0.5, colors.black),
        ("INNERGRID", (0,0), (-1,-1), 0.5, colors.black),
        ("FONTSIZE",  (0,0), (-1,-1), 8),
        ("VALIGN",    (0,0), (-1,-1), "TOP"),
        ("ROWHEIGHT", (0,1), (-1,1),  25*mm),
    ]))
    story.append(st_table)
    story.append(Spacer(1, 4*mm))

    # Note
    story.append(Paragraph(
        "<b>Note:</b> All issued assets remain the property of the Company and must be returned "
        "in good condition. Any loss or damage must be reported immediately to the Warehouse. "
        "Lost or damaged assets may be chargeable in accordance with the Company's policy.",
        small
    ))

    doc.build(story)
    return buf.getvalue()


def generate_return_pdf(data):
    from reportlab.lib.pagesizes import A4
    from reportlab.lib import colors
    from reportlab.lib.units import mm
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.enums import TA_CENTER
    import io

    buf    = io.BytesIO()
    doc    = SimpleDocTemplate(buf, pagesize=A4,
                               topMargin=10*mm, bottomMargin=10*mm,
                               leftMargin=15*mm, rightMargin=15*mm)
    styles = getSampleStyleSheet()
    story  = []

    bold   = ParagraphStyle("bold",   parent=styles["Normal"], fontName="Helvetica-Bold")
    title  = ParagraphStyle("title",  parent=styles["Normal"],
                             fontName="Helvetica-Bold", fontSize=11, alignment=TA_CENTER)
    small  = ParagraphStyle("small",  parent=styles["Normal"], fontSize=8)

    # Header with logo
    import os
    from reportlab.platypus import Image as RLImage

    logo_path = "logo.png"
    if os.path.exists(logo_path):
        logo_cell = RLImage(logo_path, width=20*mm, height=20*mm)
    else:
        logo_cell = Paragraph("", small)

    header_data = [[
        logo_cell,
        Paragraph("<b>TURCOMP ENGINEERING SERVICES SDN. BHD.</b>", title),
        Paragraph("Form No.", small),
        Paragraph("Rev. No.", small),
    ],[
        Paragraph("", small),
        Paragraph("<b>WAREHOUSE RETURN FORM</b>", title),
        Paragraph("Issue Date:", small),
        Paragraph("Page No.", small),
    ]]
    ht = Table(header_data, colWidths=[22*mm, 98*mm, 30*mm, 30*mm])
    ht.setStyle(TableStyle([
        ("BOX",       (0,0), (-1,-1), 0.5, colors.black),
        ("INNERGRID", (0,0), (-1,-1), 0.5, colors.black),
        ("VALIGN",    (0,0), (-1,-1), "MIDDLE"),
        ("ALIGN",     (0,0), (0,-1), "CENTER"),
        ("SPAN",      (0,0), (0,1)),
    ]))
    story.append(ht)
    story.append(Spacer(1, 4*mm))

    # Requestor info
    story.append(Paragraph("<b>A. REQUESTOR INFORMATION</b>", bold))
    story.append(Spacer(1, 2*mm))
    req_data = [
        ["Name", data.get("req_name",""), "Project / Job No", data.get("req_jobno","")],
        ["Project Name", data.get("req_project",""), "Usage Location", data.get("req_location","")],
        ["Date Out", data.get("date_out",""), "Date Returned", data.get("date_ret","")],
    ]
    rt = Table(req_data, colWidths=[35*mm, 65*mm, 40*mm, 40*mm])
    rt.setStyle(TableStyle([
        ("BOX",       (0,0), (-1,-1), 0.5, colors.black),
        ("INNERGRID", (0,0), (-1,-1), 0.5, colors.black),
        ("FONTNAME",  (0,0), (0,-1), "Helvetica-Bold"),
        ("FONTNAME",  (2,0), (2,-1), "Helvetica-Bold"),
        ("FONTSIZE",  (0,0), (-1,-1), 8),
        ("VALIGN",    (0,0), (-1,-1), "MIDDLE"),
    ]))
    story.append(rt)
    story.append(Spacer(1, 4*mm))

    # Asset return details
    story.append(Paragraph("<b>B. ASSET RETURN DETAILS</b>", bold))
    story.append(Spacer(1, 2*mm))
    asset_header = [["No", "Asset Tag No.", "Description", "Serial No.",
                      "Qty Returned", "Condition Out", "Condition Returned", "Status"]]
    items = data.get("items", [])
    for i, item in enumerate(items, 1):
        asset_header.append([
            str(i),
            str(item.get("TAGGING NUMBER","") or "-"),
            str(item.get("DESCRIPTION","") or "-"),
            str(item.get("SERIAL NUMBER","") or "-"),
            "1",
            str(item.get("CONDITION OUT","") or "-"),
            data.get("cond_ret",""),
            data.get("ret_status","FULLY RETURNED"),
        ])
    while len(asset_header) < 11:
        asset_header.append(["","","","","","","",""])

    at = Table(asset_header, colWidths=[10*mm, 28*mm, 42*mm, 28*mm, 15*mm, 22*mm, 22*mm, 13*mm])
    at.setStyle(TableStyle([
        ("BOX",        (0,0), (-1,-1), 0.5, colors.black),
        ("INNERGRID",  (0,0), (-1,-1), 0.5, colors.black),
        ("BACKGROUND", (0,0), (-1,0),  colors.HexColor("#1e3a5f")),
        ("TEXTCOLOR",  (0,0), (-1,0),  colors.white),
        ("FONTNAME",   (0,0), (-1,0),  "Helvetica-Bold"),
        ("FONTSIZE",   (0,0), (-1,-1), 7.5),
        ("ALIGN",      (0,0), (-1,-1), "CENTER"),
        ("VALIGN",     (0,0), (-1,-1), "MIDDLE"),
        ("ROWHEIGHT",  (0,1), (-1,-1), 8*mm),
    ]))
    story.append(at)
    story.append(Spacer(1, 4*mm))

    # Remarks
    story.append(Paragraph("<b>C. REMARKS / DISCREPANCIES</b>", bold))
    rem_table = Table([[data.get("remarks","")]], colWidths=[180*mm])
    rem_table.setStyle(TableStyle([
        ("BOX",      (0,0), (-1,-1), 0.5, colors.black),
        ("FONTSIZE", (0,0), (-1,-1), 8),
        ("MINROWHEIGHT", (0,0), (-1,-1), 20*mm),
    ]))
    story.append(rem_table)
    story.append(Spacer(1, 4*mm))

    # Overall return status
    story.append(Paragraph("<b>G. OVERALL RETURN STATUS</b>", bold))
    status_options = ["Fully Returned","Partially Returned","Outstanding","Damaged","Lost"]
    ret_stat_raw = "".join(str(data.get("ret_status","FULLY RETURNED")).strip().upper().split())
    status_row = []
    for s in status_options:
        s_clean = "".join(s.upper().split())
        if s_clean == ret_stat_raw:
            status_row.append(f"[X] {s}")
        else:
            status_row.append(f"[  ] {s}")
    st2 = Table([status_row], colWidths=[36*mm]*5)
    st2.setStyle(TableStyle([
        ("BOX",      (0,0), (-1,-1), 0.5, colors.black),
        ("FONTSIZE", (0,0), (-1,-1), 8),
        ("ALIGN",    (0,0), (-1,-1), "CENTER"),
    ]))
    story.append(st2)
    story.append(Spacer(1, 4*mm))

    # Return sign-off
    story.append(Paragraph("<b>F. RETURN SIGN-OFF</b>", bold))
    story.append(Spacer(1, 2*mm))
    keyed_by = data.get("keyed_by", "")
    req_name = data.get("req_name", "")
    now_str  = datetime.now().strftime("%Y-%m-%d %H:%M")
    sign_data = [
        [Paragraph("<b>Returned by (Requestor)</b>", small),
         Paragraph("<b>Received by<br/>(Storekeeper)</b>", small),
         Paragraph("<b>Verified by<br/>(Material Controller/Warehouse Coordinator)</b>", small)],
        [
            Paragraph(
                f"Name: {req_name}"
                f"<br/><br/><br/><br/>"
                f"Signature: ________________"
                f"<br/><br/><br/>"
                f"Date: ___________",
                small
            ),
            Paragraph(
                f"Electronically received by: <b>{keyed_by}</b><br/>"
                f"Date: {now_str}",
                small
            ),
            Paragraph(
                (f"Electronically verified by: <b>{data.get('approved_by','')}</b><br/>Date: {now_str}"
                 if data.get("approved_by")
                 else "Name: ________________<br/><br/><br/><br/>Signature: ________________<br/><br/><br/>Date: ___________"),
                small
            ),
        ],
    ]
    st3 = Table(sign_data, colWidths=[60*mm, 60*mm, 60*mm])
    st3.setStyle(TableStyle([
        ("BOX",       (0,0), (-1,-1), 0.5, colors.black),
        ("INNERGRID", (0,0), (-1,-1), 0.5, colors.black),
        ("FONTSIZE",  (0,0), (-1,-1), 8),
        ("VALIGN",    (0,0), (-1,-1), "TOP"),
        ("ROWHEIGHT", (0,1), (-1,1),  25*mm),
    ]))
    story.append(st3)
    story.append(Spacer(1, 4*mm))

    story.append(Paragraph(
        "<b>Note:</b> All issued assets remain the property of the Company and must be returned "
        "in good condition. Any loss or damage must be reported immediately to the Warehouse.",
        small
    ))

    doc.build(story)
    return buf.getvalue()

# ── WORD DOCUMENT GENERATORS ──────────────────────────────────────────────────
def generate_issue_docx(data):
    from docx import Document
    from docx.shared import Mm, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL
    import io, os

    doc = Document()
    for section in doc.sections:
        section.top_margin = Mm(12)
        section.bottom_margin = Mm(12)
        section.left_margin = Mm(15)
        section.right_margin = Mm(15)

    # Header table with logo
    ht = doc.add_table(rows=2, cols=3)
    ht.style = "Table Grid"
    if os.path.exists("logo.png"):
        cell_logo = ht.cell(0,0).merge(ht.cell(1,0))
        p_logo = cell_logo.paragraphs[0]
        run = p_logo.add_run()
        run.add_picture("logo.png", width=Mm(18))
        p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ht.cell(0,1).text = "TURCOMP ENGINEERING SERVICES SDN. BHD."
    ht.cell(0,1).paragraphs[0].runs[0].bold = True
    ht.cell(0,1).paragraphs[0].runs[0].font.size = Pt(13)
    ht.cell(1,1).text = "WAREHOUSE ISSUE FORM"
    ht.cell(1,1).paragraphs[0].runs[0].bold = True
    ht.cell(0,2).text = "Form No.\nRev. No."
    ht.cell(1,2).text = "Issue Date:\nPage No."

    doc.add_paragraph()

    # Section A
    hA = doc.add_paragraph()
    hA.add_run("A. REQUESTOR INFORMATION").bold = True

    tA = doc.add_table(rows=5, cols=4)
    tA.style = "Table Grid"
    rows_data = [
        ("Name", data.get("req_name",""), "Contact No", data.get("req_contact","")),
        ("Position", data.get("req_position",""), "Department", data.get("req_dept","")),
        ("Project Name", data.get("req_project",""), "Project / Job No", data.get("req_jobno","")),
        ("Usage Location", data.get("req_location",""), "", ""),
        ("Date Requested", data.get("req_date",""), "Required Return Date", data.get("req_retdate","")),
    ]
    for i, (a,b,c,d) in enumerate(rows_data):
        tA.cell(i,0).text = a
        tA.cell(i,0).paragraphs[0].runs[0].bold = True
        tA.cell(i,1).text = str(b)
        tA.cell(i,2).text = c
        if c: tA.cell(i,2).paragraphs[0].runs[0].bold = True
        tA.cell(i,3).text = str(d)

    doc.add_paragraph()

    # Section B — items
    hB = doc.add_paragraph()
    hB.add_run("B. ASSET ISSUE DETAILS").bold = True

    items = data.get("items", [])
    tB = doc.add_table(rows=1+max(len(items),1), cols=8)
    tB.style = "Table Grid"
    headers = ["No","Asset Tag No.","Description","Serial No.","Qty","Date Out","Condition Out","Status"]
    for j, hdr in enumerate(headers):
        c = tB.cell(0,j)
        c.text = hdr
        c.paragraphs[0].runs[0].bold = True

    for i, item in enumerate(items, 1):
        row_cells = [
            str(i),
            str(item.get("TAGGING NUMBER","") or "-"),
            str(item.get("DESCRIPTION","") or "-"),
            str(item.get("SERIAL NUMBER","") or "-"),
            "1",
            data.get("date_out",""),
            data.get("cond_out","GOOD"),
            "Issued"
        ]
        for j, val in enumerate(row_cells):
            tB.cell(i,j).text = val

    doc.add_paragraph()

    # Section E — sign-off
    hE = doc.add_paragraph()
    hE.add_run("E. ISSUE SIGN-OFF").bold = True

    now_str = datetime.now().strftime("%Y-%m-%d %H:%M")
    tE = doc.add_table(rows=2, cols=3)
    tE.style = "Table Grid"
    tE.cell(0,0).text = "Prepared by\n(Storekeeper)"
    tE.cell(0,1).text = "Checked & approved by\n(Material Controller/Warehouse Coordinator)"
    tE.cell(0,2).text = "Received by (Requestor/Receiver)"
    for c in [tE.cell(0,0), tE.cell(0,1), tE.cell(0,2)]:
        c.paragraphs[0].runs[0].bold = True

    tE.cell(1,0).text = f"Electronically signed by: {data.get('keyed_by','')}\nDate: {now_str}"
    tE.cell(1,1).text = f"Electronically approved by: {data.get('approved_by','')}\nDate: {now_str}"
    tE.cell(1,2).text = "Name: ________________\n\nSignature: ________________\n\nDate: ___________"

    doc.add_paragraph()
    note = doc.add_paragraph()
    note.add_run("Note: ").bold = True
    note.add_run("All issued assets remain the property of the Company and must be returned in good condition. Any loss or damage must be reported immediately to the Warehouse.")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def generate_return_docx(data):
    from docx import Document
    from docx.shared import Mm, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import io, os

    doc = Document()
    for section in doc.sections:
        section.top_margin = Mm(12)
        section.bottom_margin = Mm(12)
        section.left_margin = Mm(15)
        section.right_margin = Mm(15)

    ht = doc.add_table(rows=2, cols=3)
    ht.style = "Table Grid"
    if os.path.exists("logo.png"):
        cell_logo = ht.cell(0,0).merge(ht.cell(1,0))
        p_logo = cell_logo.paragraphs[0]
        run = p_logo.add_run()
        run.add_picture("logo.png", width=Mm(18))
        p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ht.cell(0,1).text = "TURCOMP ENGINEERING SERVICES SDN. BHD."
    ht.cell(0,1).paragraphs[0].runs[0].bold = True
    ht.cell(1,1).text = "WAREHOUSE RETURN FORM"
    ht.cell(1,1).paragraphs[0].runs[0].bold = True
    ht.cell(0,2).text = "Form No.\nRev. No."
    ht.cell(1,2).text = "Issue Date:\nPage No."

    doc.add_paragraph()

    hA = doc.add_paragraph()
    hA.add_run("A. REQUESTOR INFORMATION").bold = True
    tA = doc.add_table(rows=3, cols=4)
    tA.style = "Table Grid"
    rowsA = [
        ("Name", data.get("req_name",""), "Project / Job No", data.get("req_jobno","")),
        ("Project Name", data.get("req_project",""), "Usage Location", data.get("req_location","")),
        ("Date Out", data.get("date_out",""), "Date Returned", data.get("date_ret","")),
    ]
    for i,(a,b,c,d) in enumerate(rowsA):
        tA.cell(i,0).text = a; tA.cell(i,0).paragraphs[0].runs[0].bold = True
        tA.cell(i,1).text = str(b)
        tA.cell(i,2).text = c; tA.cell(i,2).paragraphs[0].runs[0].bold = True
        tA.cell(i,3).text = str(d)

    doc.add_paragraph()
    hB = doc.add_paragraph()
    hB.add_run("B. ASSET RETURN DETAILS").bold = True

    items = data.get("items", [])
    tB = doc.add_table(rows=1+max(len(items),1), cols=8)
    tB.style = "Table Grid"
    headers = ["No","Asset Tag No.","Description","Serial No.","Qty","Condition Out","Condition Returned","Status"]
    for j,hdr in enumerate(headers):
        tB.cell(0,j).text = hdr
        tB.cell(0,j).paragraphs[0].runs[0].bold = True
    for i, item in enumerate(items, 1):
        vals = [
            str(i), str(item.get("TAGGING NUMBER","") or "-"),
            str(item.get("DESCRIPTION","") or "-"), str(item.get("SERIAL NUMBER","") or "-"),
            "1", str(item.get("CONDITION OUT","") or "-"),
            data.get("cond_ret",""), data.get("ret_status","FULLY RETURNED")
        ]
        for j, v in enumerate(vals):
            tB.cell(i,j).text = v

    doc.add_paragraph()
    hC = doc.add_paragraph()
    hC.add_run("C. REMARKS / DISCREPANCIES").bold = True
    doc.add_paragraph(data.get("remarks",""))

    doc.add_paragraph()
    hG = doc.add_paragraph()
    hG.add_run("G. OVERALL RETURN STATUS").bold = True
    status_options = ["Fully Returned","Partially Returned","Outstanding","Damaged","Lost"]
    ret_stat_raw = "".join(str(data.get("ret_status","FULLY RETURNED")).strip().upper().split())
    status_line = "   ".join([
        f"[X] {s}" if "".join(s.upper().split()) == ret_stat_raw else f"[ ] {s}"
        for s in status_options
    ])
    doc.add_paragraph(status_line)

    doc.add_paragraph()
    hF = doc.add_paragraph()
    hF.add_run("F. RETURN SIGN-OFF").bold = True

    now_str = datetime.now().strftime("%Y-%m-%d %H:%M")
    tF = doc.add_table(rows=2, cols=3)
    tF.style = "Table Grid"
    tF.cell(0,0).text = "Returned by (Requestor)"
    tF.cell(0,1).text = "Received by\n(Storekeeper)"
    tF.cell(0,2).text = "Verified by\n(Material Controller/Warehouse Coordinator)"
    for c in [tF.cell(0,0), tF.cell(0,1), tF.cell(0,2)]:
        c.paragraphs[0].runs[0].bold = True

    tF.cell(1,0).text = "Name: ________________\n\nSignature: ________________\n\nDate: ___________"
    tF.cell(1,1).text = f"Electronically received by: {data.get('keyed_by','')}\nDate: {now_str}"
    tF.cell(1,2).text = f"Electronically verified by: {data.get('approved_by','')}\nDate: {now_str}" if data.get("approved_by") else "Name: ________________\n\nSignature: ________________\n\nDate: ___________"

    doc.add_paragraph()
    note = doc.add_paragraph()
    note.add_run("Note: ").bold = True
    note.add_run("All issued assets remain the property of the Company and must be returned in good condition.")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()

# ── SIDEBAR ───────────────────────────────────────────────────────────────────
with st.sidebar:
    st.markdown("""
    <div style="padding:16px 0 24px;border-bottom:1px solid #1e3a5f;margin-bottom:20px;">
        <div style="font-size:20px;font-weight:700;color:#ffffff;letter-spacing:0.5px;">
            📦 TURCOMP
        </div>
        <div style="font-size:11px;color:#64748b;margin-top:2px;">
            Inventory Management System
        </div>
    </div>
    """, unsafe_allow_html=True)
    page = st.radio("", [
        "📊 Dashboard",
        "📋 Ledger & Export",
        "✏️ Management",
        "🔄 Loan Tracker",
        "✅ Approvals",
        "⚙️ Settings"
    ], label_visibility="collapsed")
    st.markdown(f"""
    <div style="position:fixed;bottom:20px;font-size:11px;color:#475569;">
        Last refreshed<br>{datetime.now().strftime('%b %d, %Y %I:%M %p')}
    </div>
    """, unsafe_allow_html=True)

df = load_data()

# ══════════════════════════════════════════════════════════════════════════════
# PAGE 1 — DASHBOARD
# ══════════════════════════════════════════════════════════════════════════════
if page == "📊 Dashboard":
    col_title, col_filter = st.columns([3, 1])
    with col_title:
        st.markdown("## Dashboard")
        st.caption("Overview of inventory performance")
    with col_filter:
        cats     = ["All Categories"] + sorted(df["CATEGORY"].dropna().unique().tolist())
        selected = st.selectbox("Category Filter", cats, label_visibility="collapsed")

    dff = df if selected == "All Categories" else df[df["CATEGORY"] == selected]

    total     = len(dff)
    total_qty = int(dff["QUANTITY"].sum())
    out_count = int((dff["STATUS"].fillna("").str.upper() == "OUT").sum())
    available = int((dff["STATUS"].fillna("").str.upper() == "AVAILABLE").sum())

    # ── KPI CARDS ─────────────────────────────────────────────────────────────
    k1, k2, k3, k4 = st.columns(4)
    kpis = [
        (k1, "📦", "#eff6ff", "Total Items",    f"{total:,}",     "Total rows"),
        (k2, "🗂️", "#f0fdf4", "Total Quantity", f"{total_qty:,}", "Sum of quantity"),
        (k3, "📤", "#fff7ed", "Currently Out",  f"{out_count:,}", "Items on loan"),
        (k4, "✅", "#f0fdf4", "Available",      f"{available:,}", "Items available"),
    ]
    for col, icon, bg, label, val, sub in kpis:
        with col:
            st.markdown(f"""
            <div class="kpi-card">
                <div class="kpi-icon" style="background:{bg}">{icon}</div>
                <div>
                    <div class="kpi-label">{label}</div>
                    <div class="kpi-value">{val}</div>
                    <div class="kpi-sub">{sub}</div>
                </div>
            </div>
            """, unsafe_allow_html=True)

    # ── CONDITION PILLS ────────────────────────────────────────────────────────
    st.markdown('<div class="section-title">Condition Summary</div>', unsafe_allow_html=True)
    cond_series = dff["ACTUAL CONDITION"].fillna("UNKNOWN").str.upper()
    cond_counts = cond_series.value_counts().to_dict()

    pill_cols = st.columns(7)
    for i, cond in enumerate(CONDITION_FULL_OPTIONS):
        count       = cond_counts.get(cond, 0)
        icon, bg, fg = PILL_STYLE.get(cond, ("⚫", "#f1f5f9", "#334155"))
        with pill_cols[i]:
            st.markdown(f"""
            <div class="cond-pill" style="background:{bg};color:{fg};">
                <div style="font-size:11px;margin-bottom:4px;">{cond}</div>
                <div style="font-size:22px;font-weight:700;">{count:,}</div>
            </div>
            """, unsafe_allow_html=True)

    st.markdown("---")

    # ── CHARTS ROW 1: Treemap + Bar ────────────────────────────────────────────
    cat_summary = (
        dff.groupby("CATEGORY")
        .agg(Count=("DESCRIPTION", "count"),
             Quantity=("QUANTITY", "sum"))
        .reset_index()
    )

    c_left, c_mid, c_right = st.columns([1.2, 1, 1])

    with c_left:
        st.markdown('<div class="section-title">Items by Category (Treemap)</div>', unsafe_allow_html=True)
        fig_tree = px.treemap(
            cat_summary, path=["CATEGORY"], values="Count",
            custom_data=["Count", "Quantity"],
            color="Count", color_continuous_scale="Blues"
        )
        fig_tree.update_traces(
            texttemplate="<b>%{label}</b><br>%{customdata[0]}<br>(%{percentRoot:.1%})",
            hovertemplate="<b>%{label}</b><br>Items: %{customdata[0]}<br>Qty: %{customdata[1]}<extra></extra>",
            textfont_size=12
        )
        fig_tree.update_layout(
            margin=dict(t=0,l=0,r=0,b=0), height=260,
            coloraxis_showscale=False
        )
        st.plotly_chart(fig_tree, width='stretch')

    with c_mid:
        st.markdown('<div class="section-title">Quantity per Category</div>', unsafe_allow_html=True)
        fig_bar = px.bar(
            cat_summary.sort_values("Quantity", ascending=True),
            x="Quantity", y="CATEGORY", orientation="h",
            text="Quantity", color_discrete_sequence=["#2563eb"]
        )
        fig_bar.update_traces(texttemplate="%{text:,}", textposition="outside")
        fig_bar.update_layout(
            margin=dict(t=0,l=0,r=40,b=0), height=260,
            yaxis_title="", xaxis_title="",
            plot_bgcolor="white", paper_bgcolor="white",
            xaxis=dict(showgrid=True, gridcolor="#f1f5f9"),
            yaxis=dict(tickfont=dict(size=10))
        )
        st.plotly_chart(fig_bar, width='stretch')

    with c_right:
        st.markdown('<div class="section-title">Condition Breakdown</div>', unsafe_allow_html=True)
        cond_df = (
            dff["ACTUAL CONDITION"].fillna("UNKNOWN").str.upper()
            .value_counts().reset_index()
        )
        cond_df.columns = ["Condition", "Count"]
        fig_donut = go.Figure(go.Pie(
            labels=cond_df["Condition"], values=cond_df["Count"],
            hole=0.6,
            marker_colors=[COLOR_MAP_COND.get(c, "#B4B2A9") for c in cond_df["Condition"]],
            textinfo="percent",
            hovertemplate="%{label}: %{value} items (%{percent})<extra></extra>"
        ))
        fig_donut.add_annotation(
            text=f"<b>{total:,}</b><br><span style='font-size:10px'>Total Items</span>",
            x=0.5, y=0.5, showarrow=False, font_size=14
        )
        fig_donut.update_layout(
            margin=dict(t=0,l=0,r=0,b=0), height=260,
            legend=dict(font=dict(size=10), orientation="v", x=1, y=0.5),
            showlegend=True
        )
        st.plotly_chart(fig_donut, width='stretch')

    # ── LOAN STATUS ────────────────────────────────────────────────────────────
    st.markdown('<div class="section-title">Loan Status Overview</div>', unsafe_allow_html=True)
    stat_df = (
        dff["STATUS"].fillna("NOT SET").str.upper()
        .value_counts().reset_index()
    )
    stat_df.columns = ["Status", "Count"]
    stat_color = {
        "AVAILABLE":         "#16a34a",
        "OUT":               "#d97706",
        "RETURNED":          "#2563eb",
        "UNDER MAINTENANCE": "#dc2626",
        "NOT SET":           "#94a3b8",
    }
    fig_stat = px.bar(
        stat_df, x="Status", y="Count",
        color="Status",
        color_discrete_map=stat_color,
        text="Count"
    )
    fig_stat.update_traces(textposition="outside")
    fig_stat.update_layout(
        margin=dict(t=10,l=0,r=0,b=0), height=240,
        showlegend=False, plot_bgcolor="white", paper_bgcolor="white",
        yaxis=dict(showgrid=True, gridcolor="#f1f5f9"),
        xaxis_title="", yaxis_title="Number of Items"
    )
    st.plotly_chart(fig_stat, width='stretch')

# ══════════════════════════════════════════════════════════════════════════════
# PAGE 2 — LEDGER & EXPORT
# ══════════════════════════════════════════════════════════════════════════════
elif page == "📋 Ledger & Export":
    st.markdown("## Inventory Ledger & Export")

    f1, f2, f3, f4 = st.columns([2.5, 1.5, 1.5, 1])
    with f1:
        search = st.text_input("🔍 Search Brand, Model, Serial...", label_visibility="collapsed",
                               placeholder="Search Brand, Model, Serial...")
    with f2:
        cats  = ["All"] + sorted(df["CATEGORY"].dropna().unique().tolist())
        cat_f = st.selectbox("Category", cats, label_visibility="collapsed")
    with f3:
        stat_f = st.selectbox("Status", ["All Statuses"] + STATUS_OPTIONS,
                              label_visibility="collapsed")
    with f4:
        st.markdown("<br>", unsafe_allow_html=True)

    dff = df.copy()
    if cat_f  != "All":
        dff = dff[dff["CATEGORY"] == cat_f]
    if stat_f != "All Statuses":
        dff = dff[dff["STATUS"].fillna("").str.upper() == stat_f]
    if search:
        mask = dff.apply(
            lambda r: r.astype(str).str.contains(search, case=False, na=False).any(), axis=1
        )
        dff = dff[mask]

    csv = dff.to_csv(index=False).encode("utf-8")
    f1b, f2b, f3b, f4b = st.columns([2.5, 1.5, 1.5, 1])
    with f1b:
        st.caption(f"Showing {len(dff):,} of {len(df):,} items")
    with f4b:
        st.download_button(
            "⬇️ Download CSV", data=csv,
            file_name=f"inventory_{cat_f}_{datetime.now().strftime('%Y%m%d')}.csv",
            mime="text/csv", type="primary"
        )

    def highlight_row(row):
        cond   = str(row.get("ACTUAL CONDITION","")).upper()
        status = str(row.get("STATUS","")).upper()
        if cond in ["DAMAGED","DAMAGED (NOT WORKING)","LOST"]:
            return ["background-color:#f8b4b4"] * len(row)
        if status == "OUT":
            return ["background-color:#fde68a"] * len(row)
        if cond in ["DAMAGED (WORKING)","UNDER REPAIR","SERVICE DUE"]:
            return ["background-color:#fdba74"] * len(row)
        return [""] * len(row)

    display_cols = [c for c in dff.columns if c != "NO"]
    st.dataframe(
        dff[display_cols].style.apply(highlight_row, axis=1),
        height=560, width='stretch'
    )

# ══════════════════════════════════════════════════════════════════════════════
# PAGE 3 — MANAGEMENT
# ══════════════════════════════════════════════════════════════════════════════
elif page == "✏️ Management":
    st.markdown("## Inventory Management")
    ws   = get_worksheet()
    tab1, tab2, tab3 = st.tabs(["➕ Add New Item", "🔧 Edit Existing Item", "🗑️ Delete Item"])

    # ── TAB 1: ADD ─────────────────────────────────────────────────────────────
    with tab1:
        with st.form("add_form"):
            st.markdown("#### New Item Details")
            c1, c2 = st.columns(2)
            with c1:
                existing_cats = load_categories()
                cat_choice = st.selectbox("Category", ["-- Select --"] + existing_cats)
                category = "" if cat_choice == "-- Select --" else cat_choice
                st.caption("💡 Need a new category? Go to ⚙️ Settings to add one.")
                type_spec   = st.text_input("Type/Spec")
                brand       = st.text_input("Brand")
                model       = st.text_input("Model")
                serial      = st.text_input("Serial Number", placeholder="Optional")
                tagging     = st.text_input("Tagging Number")
                description = st.text_input("Description ✱", placeholder="Required")
            with c2:
                size        = st.text_input("Size")
                watt        = st.text_input("Watt")
                thickness   = st.text_input("Thickness")
                quantity    = st.number_input("Quantity", min_value=0, value=1)
                uom         = st.text_input("UOM", value="UNIT")
                storage_loc = st.text_input("Storage Location")
                condition   = st.selectbox("Actual Condition", CONDITION_FULL_OPTIONS)
                remarks     = st.text_input("Remarks")

            st.info("ℹ️ New items will be automatically set to STATUS = AVAILABLE")

            if st.form_submit_button("➕ Add Item", type="primary"):
                if not description:
                    st.error("❌ Description is required.")
                elif serial and serial in df["SERIAL NUMBER"].fillna("").astype(str).values:
                    st.error(f"❌ Serial Number '{serial}' already exists.")
                else:
                    next_row = len(df) + 2  # +2 because row 1 is headers
                    new_data = [
                        "", category, type_spec, brand, model,
                        serial, tagging, description, size, watt, thickness,
                        quantity, uom, storage_loc, condition,
                        "", "", "", "", "", "", "", "", "AVAILABLE", remarks
                    ]

                    # Auto-expand the sheet if it's running out of rows
                    current_row_count = ws.row_count
                    if next_row > current_row_count:
                        ws.add_rows(200)  # add 200 extra rows as buffer

                    ws.update(f"A{next_row}:Y{next_row}", [new_data])
                    st.success(f"✅ Item '{description}' added successfully at row {next_row}! ({datetime.now().strftime('%Y-%m-%d %H:%M')})")
                    reload()

    # ── TAB 2: EDIT ─────────────────────────────────────────────────────────────
    with tab2:
        st.markdown("#### Find & Edit Item")
        row, row_num = search_item(df, label="edit")

        if row is not None:
            from datetime import datetime as dt

            def parse_date_safe(val):
                val = str(val or "").strip()
                if not val or val in ["nan", "None"]:
                    return None
                for fmt in ["%Y-%m-%d", "%Y/%m/%d", "%d/%m/%Y", "%d-%m-%Y"]:
                    try:
                        return dt.strptime(val, fmt).date()
                    except:
                        continue
                return None

            st.caption("💡 Tip: Date Out is only needed when Status = OUT. Date Returned is only needed after it comes back. Leave the checkbox unticked to keep a date empty.")
            dc1, dc2 = st.columns(2)
            with dc1:
                date_out_has_value = st.checkbox(
                    "This item has a Date Out (e.g. Status is OUT or was borrowed before)",
                    value=parse_date_safe(row.get("DATE OUT")) is not None,
                    key="chk_date_out"
                )
            with dc2:
                date_ret_has_value = st.checkbox(
                    "This item has a Date Returned (e.g. it has already come back)",
                    value=parse_date_safe(row.get("DATE RETURNED")) is not None,
                    key="chk_date_ret"
                )

            with st.form("edit_form"):
                st.markdown("**Inventory Details**")
                c1, c2 = st.columns(2)
                with c1:
                    existing_cats_edit = load_categories()
                    current_cat = str(row.get("CATEGORY") or "")
                    cat_options = existing_cats_edit if current_cat in existing_cats_edit else existing_cats_edit + [current_cat]
                    cat_idx = cat_options.index(current_cat) if current_cat in cat_options else 0
                    e_cat = st.selectbox("Category", cat_options, index=cat_idx)
                    st.caption("💡 Need a new category? Go to ⚙️ Settings to add one.")
                    e_type   = st.text_input("Type/Spec",       value=str(row.get("TYPE/SPEC")    or ""))
                    e_brand  = st.text_input("Brand",           value=str(row.get("BRAND")        or ""))
                    e_model  = st.text_input("Model",           value=str(row.get("MODEL")        or ""))
                    e_serial = st.text_input("Serial Number",   value=str(row.get("SERIAL NUMBER")or ""))
                    e_tag    = st.text_input("Tagging Number",  value=str(row.get("TAGGING NUMBER")or ""))
                    e_desc   = st.text_input("Description",     value=str(row.get("DESCRIPTION")  or ""))
                with c2:
                    e_size   = st.text_input("Size",            value=str(row.get("SIZE")         or ""))
                    e_watt   = st.text_input("Watt",            value=str(row.get("WATT")         or ""))
                    e_thick  = st.text_input("Thickness",       value=str(row.get("THICKNESS")    or ""))
                    e_qty    = st.number_input("Quantity",      value=int(row.get("QUANTITY")     or 0))
                    e_uom    = st.text_input("UOM",             value=str(row.get("UOM")          or ""))
                    e_stor   = st.text_input("Storage Location",value=str(row.get("STORAGE LOCATION") or ""))
                    cur_ac   = str(row.get("ACTUAL CONDITION","GOOD")).upper()
                    ac_idx   = CONDITION_FULL_OPTIONS.index(cur_ac) if cur_ac in CONDITION_FULL_OPTIONS else 0
                    e_cond   = st.selectbox("Actual Condition", CONDITION_FULL_OPTIONS, index=ac_idx)
                    e_rem    = st.text_input("Remarks",         value=str(row.get("REMARKS")      or ""))

                st.markdown("---")
                st.markdown("**Loan Status (Admin Override)**")
                a1, a2 = st.columns(2)
                with a1:
                    cur_stat = str(row.get("STATUS","AVAILABLE")).upper()
                    st._idx  = STATUS_OPTIONS.index(cur_stat) if cur_stat in STATUS_OPTIONS else 0
                    e_status = st.selectbox("Status", STATUS_OPTIONS, index=st._idx)
                    from datetime import datetime as dt

                    def parse_date_safe(val):
                        """Try to parse existing date string, return None if empty/invalid"""
                        val = str(val or "").strip()
                        if not val or val in ["nan", "None"]:
                            return None
                        for fmt in ["%Y-%m-%d", "%Y/%m/%d", "%d/%m/%Y", "%d-%m-%Y"]:
                            try:
                                return dt.strptime(val, fmt).date()
                            except:
                                continue
                        return None

                    if date_out_has_value:
                        e_dateout_val = st.date_input(
                            "Date Out",
                            value=parse_date_safe(row.get("DATE OUT")) or dt.today().date()
                        )
                        e_dateout = str(e_dateout_val)
                    else:
                        st.caption("Date Out: (not set)")
                        e_dateout = ""

                    if date_ret_has_value:
                        e_dateret_val = st.date_input(
                            "Date Returned",
                            value=parse_date_safe(row.get("DATE RETURNED")) or dt.today().date()
                        )
                        e_dateret = str(e_dateret_val)
                    else:
                        st.caption("Date Returned: (not set)")
                        e_dateret = ""
                    e_jobno  = st.text_input("Job No",         value=str(row.get("JOB NO")        or ""))
                with a2:
                    e_req    = st.text_input("Requestor",      value=str(row.get("REQUESTOR")     or ""))
                    e_proj   = st.text_input("Project/Usage",  value=str(row.get("PROJECT / USAGE") or ""))
                    e_loc    = st.text_input("Location",       value=str(row.get("LOCATION")      or ""))
                    cur_co   = str(row.get("CONDITION OUT","GOOD")).upper()
                    co_idx   = CONDITION_OUT_OPTIONS.index(cur_co) if cur_co in CONDITION_OUT_OPTIONS else 0
                    e_co     = st.selectbox("Condition Out", CONDITION_OUT_OPTIONS, index=co_idx)
                    condition_returned_options = ["N/A (Not Returned Yet)"] + CONDITION_FULL_OPTIONS
                    cur_cr   = str(row.get("CONDITION RETURNED","") or "").upper()
                    if not cur_cr or cur_cr in ["NAN", "NONE", ""]:
                        cr_idx = 0
                    else:
                        cr_idx = condition_returned_options.index(cur_cr) if cur_cr in condition_returned_options else 0
                    e_cr_choice = st.selectbox("Condition Returned", condition_returned_options, index=cr_idx)
                    e_cr = "" if e_cr_choice == "N/A (Not Returned Yet)" else e_cr_choice

                if st.form_submit_button("💾 Save All Changes", type="primary"):
                    try:
                        updated = [
                            "", e_cat, e_type, e_brand, e_model,
                            e_serial, e_tag, e_desc, e_size, e_watt, e_thick,
                            e_qty, e_uom, e_stor, e_cond,
                            e_dateout, e_dateret, e_req, e_jobno, e_proj,
                            e_loc, e_co, e_cr, e_status, e_rem
                        ]
                        ws.update(f"A{row_num}:Y{row_num}", [updated])
                        st.success(f"✅ Item '{e_desc}' has been updated successfully! ({datetime.now().strftime('%Y-%m-%d %H:%M:%S')})")
                        reload()
                    except Exception as e:
                        st.error(f"❌ Failed to save changes. Error: {str(e)}")

    # ── TAB 3: DELETE ─────────────────────────────────────────────────────────
    with tab3:
        st.markdown("#### Delete Item")
        st.warning("⚠️ Deleted rows are **permanently removed** from the Google Sheet and cannot be recovered.")
        row, row_num = search_item(df, label="delete")

        if row is not None:
            st.error(f"""
            **You are about to permanently delete this item:**

            - Description: {row.get('DESCRIPTION','—')}
            - Serial Number: {row.get('SERIAL NUMBER','—')}
            - Tagging Number: {row.get('TAGGING NUMBER','—')}
            - Category: {row.get('CATEGORY','—')}
            - Current Status: {row.get('STATUS','—')}
            """)

            confirm = st.checkbox("I understand this action is permanent and cannot be undone")
            if confirm:
                if st.button("🗑️ Permanently Delete This Item", type="primary"):
                    ws.delete_rows(int(row_num))
                    st.success(f"✅ Item '{row.get('DESCRIPTION','')}' has been permanently deleted from the system.")
                    reload()

# ══════════════════════════════════════════════════════════════════════════════
# PAGE 4 — LOAN TRACKER
# ══════════════════════════════════════════════════════════════════════════════
elif page == "🔄 Loan Tracker":
    st.markdown("## 🔄 Loan Tracker")
    ws = get_worksheet()
    tab1, tab2, tab3 = st.tabs([
        "🛒 Issue Items Out",
        "✅ Return Item",
        "📋 Currently Out"
    ])

    # ── TAB 1: BASKET ISSUE OUT ───────────────────────────────────────────────
    with tab1:
        # Init basket
        if "basket" not in st.session_state:
            st.session_state.basket = []

        st.markdown("#### Step 1 — Requestor Information")
        ri1, ri2 = st.columns(2)
        with ri1:
            req_name     = st.text_input("Name ✱", key="req_name")
            req_position = st.text_input("Position", key="req_pos")
            req_project  = st.text_input("Project Name", key="req_proj")
            req_location = st.text_input("Usage Location", key="req_loc")
        with ri2:
            req_contact  = st.text_input("Contact No", key="req_contact")
            req_dept     = st.text_input("Department", key="req_dept")
            req_jobno    = st.text_input("Project / Job No ✱", key="req_jobno")
            req_date     = st.date_input("Date Requested", value=datetime.today())
            req_retdate  = st.date_input("Required Return Date")

        st.markdown("---")
        st.markdown("#### Step 2 — Search & Add Items to Basket")
        st.caption("Filter by category and search first, then add items to your basket.")

        sc1, sc2 = st.columns(2)
        with sc1:
            filter_cat = st.selectbox(
                "Filter by Category ✱",
                ["-- Select Category --"] + sorted(df["CATEGORY"].dropna().unique().tolist()),
                key="basket_cat"
            )
        with sc2:
            filter_search = st.text_input(
                "Search Description / Serial / Tagging (optional)",
                key="basket_search"
            )

        # Only show items when category is selected
        if filter_cat == "-- Select Category --":
            st.info("👆 Please select a category above to see available items.")
        else:
            filtered = df[
                (df["CATEGORY"] == filter_cat) &
                (df["STATUS"].fillna("").str.upper().isin(["AVAILABLE", ""]))
            ]
            if filter_search:
                mask = filtered.apply(
                    lambda r: r.astype(str).str.contains(
                        filter_search, case=False, na=False).any(), axis=1
                )
                filtered = filtered[mask]

            # Remove already-in-basket items
            basket_keys = [b["_idx"] for b in st.session_state.basket]
            filtered    = filtered[~filtered.index.isin(basket_keys)]

            if filtered.empty:
                st.warning("No available items found for this filter.")
            else:
                st.caption(f"{len(filtered)} available items")
                show_cols = [c for c in [
                    "TAGGING NUMBER","DESCRIPTION","SERIAL NUMBER",
                    "BRAND","SIZE","ACTUAL CONDITION"
                ] if c in filtered.columns]

                # Header row
                hcols = st.columns([0.4] + [1]*len(show_cols) + [0.6])
                hcols[0].markdown("**Add**")
                for i, c in enumerate(show_cols):
                    hcols[i+1].markdown(f"**{c}**")

                st.markdown("<hr style='margin:3px 0'>", unsafe_allow_html=True)

                for idx, row_data in filtered.head(50).iterrows():
                    rcols = st.columns([0.4] + [1]*len(show_cols) + [0.6])
                    for i, c in enumerate(show_cols):
                        val = str(row_data.get(c,"") or "")
                        rcols[i+1].caption(val if val not in ["None","nan","","-"] else "-")
                    if rcols[-1].button("➕ Add", key=f"add_{idx}"):
                        item_dict = row_data.to_dict()
                        item_dict["_idx"] = idx
                        st.session_state.basket.append(item_dict)
                        st.rerun()

                if len(filtered) > 50:
                    st.caption(f"Showing first 50 of {len(filtered)}. Refine your search to see more.")

        st.markdown("---")

        # ── BASKET DISPLAY ────────────────────────────────────────────────────
        st.markdown(f"#### 🛒 Basket ({len(st.session_state.basket)} items)")
        if not st.session_state.basket:
            st.info("Your basket is empty. Add items above.")
        else:
            basket_cols = ["TAGGING NUMBER","DESCRIPTION","SERIAL NUMBER","BRAND","SIZE","ACTUAL CONDITION"]
            basket_df   = pd.DataFrame(st.session_state.basket)
            disp_cols   = [c for c in basket_cols if c in basket_df.columns]
            st.dataframe(basket_df[disp_cols], use_container_width=True)

            # Remove item from basket
            remove_options = [
                f"{i+1}. {b.get('DESCRIPTION','')} | {b.get('TAGGING NUMBER','')}"
                for i, b in enumerate(st.session_state.basket)
            ]
            rem_col1, rem_col2 = st.columns([3,1])
            with rem_col1:
                to_remove = st.selectbox("Remove an item from basket", ["-- Keep all --"] + remove_options)
            with rem_col2:
                st.markdown("<br>", unsafe_allow_html=True)
                if st.button("🗑️ Remove") and to_remove != "-- Keep all --":
                    idx_to_remove = remove_options.index(to_remove)
                    st.session_state.basket.pop(idx_to_remove)
                    st.rerun()

            if st.button("🗑️ Clear Entire Basket"):
                st.session_state.basket = []
                st.rerun()

        st.markdown("---")

        # ── STEP 3: APPROVAL ──────────────────────────────────────────────────
        if st.session_state.basket:
            st.markdown("#### Step 3 — Issue Details & Approval")

            ic1, ic2 = st.columns(2)
            with ic1:
                date_out  = st.date_input("Date Out", value=datetime.today(), key="bulk_date")
                cond_out  = st.selectbox("Condition Out", CONDITION_OUT_OPTIONS, key="bulk_cond")
                keyed_by  = st.text_input("Keyed in by ✱", placeholder="Enter your name")
            with ic2:
                st.markdown("**Approval Action**")
                approval_choice = st.radio(
                    "Choose action",
                    ["✅ Self-Approve (I am authorised to approve)",
                     "⏳ Send to Waitlist (needs another person to approve)"],
                    key="approval_choice"
                )
                if "Self-Approve" in approval_choice:
                    st.success("You will approve this request immediately. PDF can be generated right after.")
                else:
                    st.warning("Request will go to the Approval Queue. PDF can only be printed after approval.")

            if st.button("🚀 Submit Request", type="primary"):
                if not req_name:
                    st.error("❌ Requestor name is required.")
                elif not req_jobno:
                    st.error("❌ Job No is required.")
                elif not keyed_by:
                    st.error("❌ Please enter who is keying in this request.")
                else:
                    approval_status = "APPROVED" if "Self-Approve" in approval_choice else "PENDING"
                    approved_by     = keyed_by if "Self-Approve" in approval_choice else ""
                    request_id      = str(uuid.uuid4())[:8].upper()
                    items_json      = json.dumps([{
                        "TAGGING NUMBER":  b.get("TAGGING NUMBER",""),
                        "DESCRIPTION":     b.get("DESCRIPTION",""),
                        "SERIAL NUMBER":   b.get("SERIAL NUMBER",""),
                        "BRAND":           b.get("BRAND",""),
                        "SIZE":            b.get("SIZE",""),
                        "ACTUAL CONDITION":b.get("ACTUAL CONDITION",""),
                        "_idx":            b.get("_idx"),
                    } for b in st.session_state.basket])

                    # Save to approvals sheet
                    aws = get_approvals_sheet()
                    aws.append_row([
                        request_id, keyed_by,
                        datetime.now().strftime("%Y-%m-%d %H:%M"),
                        req_name, req_contact, req_position, req_dept,
                        req_project, req_jobno, req_location,
                        str(req_date), str(req_retdate),
                        approval_status, approved_by, items_json
                    ])

                    # If approved — mark items OUT immediately
                    if approval_status == "APPROVED":
                        for b in st.session_state.basket:
                            row_num = b["_idx"] + 2
                            try:
                                ws.update(f"P{row_num}:Y{row_num}", [[
                                    str(date_out), "", req_name,
                                    req_jobno, req_project, req_location,
                                    cond_out, "", "OUT",
                                    str(b.get("REMARKS","") or "")
                                ]])
                            except:
                                pass

                    st.session_state.basket = []
                    st.cache_data.clear()

                    if approval_status == "APPROVED":
                        st.success(f"✅ Request **{request_id}** approved by **{keyed_by}**! Items marked OUT. You can now print the PDF from the Approvals page.")
                    else:
                        st.warning(f"⏳ Request **{request_id}** sent to waitlist. Waiting for approval before PDF can be printed.")
                    st.rerun()

 # ── TAB 2: RETURN ─────────────────────────────────────────────────────────
# ── TAB 2: BASKET RETURN ──────────────────────────────────────────────────
    with tab2:
        if "return_basket" not in st.session_state:
            st.session_state.return_basket = []

        st.markdown("#### Step 1 — Search & Add Items to Return")
        st.caption("Search for items currently OUT and add them to your return basket.")

        rc1, rc2 = st.columns(2)
        with rc1:
            ret_filter_cat = st.selectbox(
                "Filter by Category",
                ["All"] + sorted(df["CATEGORY"].dropna().unique().tolist()),
                key="ret_basket_cat"
            )
        with rc2:
            ret_filter_search = st.text_input(
                "Search Description / Serial / Tagging / Requestor",
                key="ret_basket_search"
            )

        out_items = df[df["STATUS"].fillna("").str.upper().isin(["OUT","PARTIALLY RETURNED"])]
        if ret_filter_cat != "All":
            out_items = out_items[out_items["CATEGORY"] == ret_filter_cat]
        if ret_filter_search:
            mask = out_items.apply(
                lambda r: r.astype(str).str.contains(ret_filter_search, case=False, na=False).any(), axis=1
            )
            out_items = out_items[mask]

        ret_basket_keys = [b["_idx"] for b in st.session_state.return_basket]
        out_items = out_items[~out_items.index.isin(ret_basket_keys)]

        if out_items.empty:
            st.info("No matching items currently out on loan.")
        else:
            st.caption(f"{len(out_items)} items currently out")
            show_cols = [c for c in [
                "TAGGING NUMBER","DESCRIPTION","SERIAL NUMBER",
                "REQUESTOR","JOB NO","DATE OUT"
            ] if c in out_items.columns]

            hcols = st.columns([0.4] + [1]*len(show_cols))
            hcols[0].markdown("**Add**")
            for i, c in enumerate(show_cols):
                hcols[i+1].markdown(f"**{c}**")
            st.markdown("<hr style='margin:3px 0'>", unsafe_allow_html=True)

            for idx, row_data in out_items.head(50).iterrows():
                rcols = st.columns([0.4] + [1]*len(show_cols))
                for i, c in enumerate(show_cols):
                    val = str(row_data.get(c,"") or "")
                    rcols[i+1].caption(val if val not in ["None","nan","",""] else "-")
                if rcols[0].button("➕", key=f"ret_add_{idx}"):
                    item_dict = row_data.to_dict()
                    item_dict["_idx"] = idx
                    st.session_state.return_basket.append(item_dict)
                    st.rerun()

            if len(out_items) > 50:
                st.caption(f"Showing first 50 of {len(out_items)}. Refine your search to see more.")

        st.markdown("---")
        st.markdown(f"#### 🛒 Return Basket ({len(st.session_state.return_basket)} items)")

        if not st.session_state.return_basket:
            st.info("Your return basket is empty. Add items above.")
        else:
            basket_cols = ["TAGGING NUMBER","DESCRIPTION","SERIAL NUMBER","REQUESTOR","JOB NO","DATE OUT"]
            basket_df   = pd.DataFrame(st.session_state.return_basket)
            disp_cols   = [c for c in basket_cols if c in basket_df.columns]
            st.dataframe(basket_df[disp_cols], use_container_width=True)

            remove_options = [
                f"{i+1}. {b.get('DESCRIPTION','')} | {b.get('TAGGING NUMBER','')}"
                for i, b in enumerate(st.session_state.return_basket)
            ]
            rem_col1, rem_col2 = st.columns([3,1])
            with rem_col1:
                to_remove = st.selectbox("Remove an item", ["-- Keep all --"] + remove_options, key="ret_remove_sel")
            with rem_col2:
                st.markdown("<br>", unsafe_allow_html=True)
                if st.button("🗑️ Remove", key="ret_remove_btn") and to_remove != "-- Keep all --":
                    idx_to_remove = remove_options.index(to_remove)
                    st.session_state.return_basket.pop(idx_to_remove)
                    st.rerun()

            if st.button("🗑️ Clear Entire Return Basket", key="ret_clear_all"):
                st.session_state.return_basket = []
                st.rerun()

        st.markdown("---")

        if st.session_state.return_basket:
            st.markdown("#### Step 2 — Return Details")

            def clean(v):
                s = str(v) if v is not None else ""
                return "" if s in ["nan","None","NaN","-"] else s

            rc1, rc2 = st.columns(2)
            with rc1:
                date_ret   = st.date_input("Date Returned", value=datetime.today(), key="basket_date_ret")
                ret_status = st.selectbox("Return Status (applies to all)", RETURN_STATUS_OPTIONS, key="basket_ret_status")
                keyed_by   = st.text_input("Processed by (your name) ✱", key="basket_keyed_by")
            with rc2:
                cond_ret   = st.selectbox("Condition Returned (applies to all)", CONDITION_FULL_OPTIONS, key="basket_cond_ret")
                new_cond   = st.selectbox("Update Actual Condition (applies to all)", CONDITION_FULL_OPTIONS, key="basket_new_cond")
                remarks    = st.text_input("Remarks", key="basket_ret_remarks")

            status_map = {
                "FULLY RETURNED":     "AVAILABLE",
                "PARTIALLY RETURNED": "PARTIALLY RETURNED",
                "LOST":               "LOST",
                "DAMAGED":            "AVAILABLE",
                "OUTSTANDING":        "OUT",
            }
            new_status = status_map.get(ret_status, "AVAILABLE")

            st.markdown("**Approval**")
            ret_approval_choice = st.radio(
                "Choose action",
                ["✅ Self-Approve (I am authorised to approve)",
                 "⏳ Send to Waitlist (needs another person to approve)"],
                key="ret_approval_choice"
            )

            b1, b2, b3 = st.columns(3)

            with b1:
                if st.button("🚀 Submit Return Request", type="primary", key="basket_confirm_ret"):
                    if not keyed_by:
                        st.error("❌ Please enter who is processing this return.")
                    else:
                        approval_status = "APPROVED" if "Self-Approve" in ret_approval_choice else "PENDING"
                        approved_by     = keyed_by if "Self-Approve" in ret_approval_choice else ""
                        request_id      = str(uuid.uuid4())[:8].upper()
                        items_json      = json.dumps([{
                            "TAGGING NUMBER": b.get("TAGGING NUMBER",""),
                            "DESCRIPTION":    b.get("DESCRIPTION",""),
                            "SERIAL NUMBER":  b.get("SERIAL NUMBER",""),
                            "REQUESTOR":      b.get("REQUESTOR",""),
                            "JOB NO":         b.get("JOB NO",""),
                            "DATE OUT":       b.get("DATE OUT",""),
                            "CONDITION OUT":  b.get("CONDITION OUT",""),
                            "PROJECT / USAGE":b.get("PROJECT / USAGE",""),
                            "LOCATION":       b.get("LOCATION",""),
                            "_idx":           b.get("_idx"),
                        } for b in st.session_state.return_basket])

                        aws = get_approvals_sheet()
                        first_basket_item = st.session_state.return_basket[0] if st.session_state.return_basket else {}
                        aws.append_row([
                            request_id, keyed_by,
                            datetime.now().strftime("%Y-%m-%d %H:%M"),
                            "Multiple" if len(st.session_state.return_basket) > 1 else clean(first_basket_item.get("REQUESTOR","")),
                            "", "", "",
                            clean(first_basket_item.get("PROJECT / USAGE","")),
                            clean(first_basket_item.get("JOB NO","")),
                            clean(first_basket_item.get("LOCATION","")),
                            str(date_ret), str(date_ret),
                            approval_status, approved_by, items_json, "RETURN",
                            "", "", ret_status
                        ])

                        if approval_status == "APPROVED":
                            for b in st.session_state.return_basket:
                                row_num = b["_idx"] + 2
                                try:
                                    ws.update(f"P{row_num}:Y{row_num}", [[
                                        clean(b.get("DATE OUT","")), str(date_ret),
                                        clean(b.get("REQUESTOR","")),
                                        clean(b.get("JOB NO","")),
                                        clean(b.get("PROJECT / USAGE","")),
                                        clean(b.get("LOCATION","")),
                                        clean(b.get("CONDITION OUT","")),
                                        cond_ret, new_status, remarks
                                    ]])
                                    ws.update(f"O{row_num}", [[new_cond]])
                                except:
                                    pass

                        first_basket_item2 = st.session_state.return_basket[0] if st.session_state.return_basket else {}
                        st.session_state["last_return_batch"] = {
                            "keyed_by":     keyed_by, "approved_by": approved_by,
                            "req_jobno":    clean(first_basket_item2.get("JOB NO","")),
                            "req_project":  clean(first_basket_item2.get("PROJECT / USAGE","")),
                            "req_location": clean(first_basket_item2.get("LOCATION","")),
                            "date_out":     clean(first_basket_item2.get("DATE OUT","")),
                            "date_ret":     str(date_ret),
                            "ret_status":   ret_status, "cond_ret": cond_ret,
                            "remarks":      remarks, "items": st.session_state.return_basket.copy(),
                            "req_name":     "Multiple" if len(st.session_state.return_basket) > 1 else clean(first_basket_item2.get("REQUESTOR","")),
                        }
                        st.session_state.return_basket = []
                        st.cache_data.clear()

                        if approval_status == "APPROVED":
                            st.success(f"✅ Request {request_id} approved by {keyed_by}! Items returned. PDF/Word available below.")
                        else:
                            st.warning(f"⏳ Request {request_id} sent to waitlist for approval.")
                        st.rerun()

            with b2:
                if keyed_by and st.session_state.return_basket:
                    first_item = st.session_state.return_basket[0] if st.session_state.return_basket else {}
                    preview_data = {
                        "req_name":     "Multiple" if len(st.session_state.return_basket) > 1 else clean(first_item.get("REQUESTOR","")),
                        "req_jobno":    clean(first_item.get("JOB NO","")),
                        "req_project":  clean(first_item.get("PROJECT / USAGE","")),
                        "req_location": clean(first_item.get("LOCATION","")),
                        "date_out":     clean(first_item.get("DATE OUT","")),
                        "date_ret":     str(date_ret), "ret_status": ret_status,
                        "cond_ret":     cond_ret, "remarks": remarks, "keyed_by": keyed_by,
                        "items":        st.session_state.return_basket
                    }
                    st.download_button("📄 Preview PDF", data=generate_return_pdf(preview_data),
                        file_name=f"Return_{keyed_by}_{date_ret}.pdf", mime="application/pdf", key="basket_ret_pdf_prev")

            with b3:
                if keyed_by and st.session_state.return_basket:
                    first_item = st.session_state.return_basket[0] if st.session_state.return_basket else {}
                    preview_data = {
                        "req_name":     "Multiple" if len(st.session_state.return_basket) > 1 else clean(first_item.get("REQUESTOR","")),
                        "req_jobno":    clean(first_item.get("JOB NO","")),
                        "req_project":  clean(first_item.get("PROJECT / USAGE","")),
                        "req_location": clean(first_item.get("LOCATION","")),
                        "date_out":     clean(first_item.get("DATE OUT","")),
                        "date_ret":     str(date_ret), "ret_status": ret_status,
                        "cond_ret":     cond_ret, "remarks": remarks, "keyed_by": keyed_by,
                        "items":        st.session_state.return_basket
                    }
                    st.download_button("📝 Preview Word", data=generate_return_docx(preview_data),
                        file_name=f"Return_{keyed_by}_{date_ret}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        key="basket_ret_docx_prev")

        # Show last confirmed batch PDF if just returned
        if "last_return_batch" in st.session_state and st.session_state["last_return_batch"]:
            batch = st.session_state["last_return_batch"]
            st.markdown("---")
            st.markdown("#### ✅ Last Return Completed")
            pdf_bytes = generate_return_pdf(batch)
            st.download_button(
                "📄 Download Return Form PDF (Last Batch)",
                data=pdf_bytes,
                file_name=f"Return_{batch.get('keyed_by','')}_{batch.get('date_ret','')}.pdf",
                mime="application/pdf",
                key="last_batch_pdf"
            )

    # ── TAB 3: CURRENTLY OUT ──────────────────────────────────────────────────
    with tab3:
        out_df = df[df["STATUS"].fillna("").str.upper().isin(["OUT","PARTIALLY RETURNED"])]
        if out_df.empty:
            st.info("✅ No items currently out on loan.")
        else:
            st.markdown(f"#### {len(out_df)} Items Currently Out")
            show_cols = [c for c in [
                "TAGGING NUMBER","DESCRIPTION","SERIAL NUMBER","CATEGORY",
                "BRAND","REQUESTOR","JOB NO","DATE OUT",
                "PROJECT / USAGE","LOCATION","CONDITION OUT","STATUS"
            ] if c in out_df.columns]
            st.dataframe(out_df[show_cols], height=500, width='stretch')


# ══════════════════════════════════════════════════════════════════════════════
# PAGE 5 — APPROVALS
# ══════════════════════════════════════════════════════════════════════════════
elif page == "✅ Approvals":
    st.markdown("## ✅ Approvals")
    archive_old_approvals()
    aws  = get_approvals_sheet()
    adf  = load_approvals()

    approval_search = st.text_input(
        "🔍 Search all approvals (Request ID, Name, Job No, Keyed by)",
        key="approval_search"
    )
    if approval_search and not adf.empty:
        mask = adf.apply(
            lambda r: r.astype(str).str.contains(approval_search, case=False, na=False).any(), axis=1
        )
        adf = adf[mask]

    tab_pending, tab_approved = st.tabs(["⏳ Pending Approval", "✅ Approved Requests"])

    # ── PENDING ───────────────────────────────────────────────────────────────
    with tab_pending:
        pending = adf[adf["STATUS"] == "PENDING"] if not adf.empty else pd.DataFrame()
        if pending.empty:
            st.success("✅ No requests waiting for approval.")
        else:
            st.markdown(f"#### {len(pending)} Request(s) Waiting for Approval")
            for _, req in pending.iterrows():
                req_type = req.get("TYPE", "ISSUE") or "ISSUE"
                type_icon = "🚚" if req_type == "ISSUE" else "↩️"
                with st.expander(
                    f"{type_icon} [{req_type}] Request {req['REQUEST_ID']} — {req['REQ_NAME']} | "
                    f"Job: {req['REQ_JOBNO']} | Keyed by: {req['KEYED_BY']} | "
                    f"{req['DATE_REQUESTED']}"
                ):
                    st.markdown(f"""
                    **Requestor:** {req['REQ_NAME']} | **Contact:** {req['REQ_CONTACT']}
                    **Position:** {req['REQ_POSITION']} | **Department:** {req['REQ_DEPT']}
                    **Project:** {req['REQ_PROJECT']} | **Job No:** {req['REQ_JOBNO']}
                    **Location:** {req['REQ_LOCATION']}
                    **Date Requested:** {req['REQ_DATE']} | **Return Date:** {req['REQ_RETDATE']}
                    """)

                    # Show items
                    try:
                        items = json.loads(req["ITEMS_JSON"])
                        items_df = pd.DataFrame(items)
                        disp = [c for c in ["TAGGING NUMBER","DESCRIPTION","SERIAL NUMBER","BRAND","SIZE"] if c in items_df.columns]
                        st.dataframe(items_df[disp], use_container_width=True)
                        with st.expander(f"✏️ Edit Request Details — {req['REQUEST_ID']}"):
                            ec1, ec2 = st.columns(2)
                            with ec1:
                                edit_name    = st.text_input("Requestor Name", value=req.get("REQ_NAME","") or "", key=f"edit_name_{req['REQUEST_ID']}")
                                edit_contact = st.text_input("Contact No",     value=req.get("REQ_CONTACT","") or "", key=f"edit_contact_{req['REQUEST_ID']}")
                                edit_pos     = st.text_input("Position",      value=req.get("REQ_POSITION","") or "", key=f"edit_pos_{req['REQUEST_ID']}")
                                edit_dept    = st.text_input("Department",    value=req.get("REQ_DEPT","") or "", key=f"edit_dept_{req['REQUEST_ID']}")
                            with ec2:
                                edit_proj    = st.text_input("Project",       value=req.get("REQ_PROJECT","") or "", key=f"edit_proj_{req['REQUEST_ID']}")
                                edit_job     = st.text_input("Job No",        value=req.get("REQ_JOBNO","") or "", key=f"edit_job_{req['REQUEST_ID']}")
                                edit_loc     = st.text_input("Location",      value=req.get("REQ_LOCATION","") or "", key=f"edit_loc_{req['REQUEST_ID']}")
                                edit_keyed   = st.text_input("Keyed By",      value=req.get("KEYED_BY","") or "", key=f"edit_keyed_{req['REQUEST_ID']}")

                            if st.button(f"💾 Save Edits — {req['REQUEST_ID']}", key=f"save_edit_{req['REQUEST_ID']}"):
                                all_reqs = aws.get_all_values()
                                for i, r in enumerate(all_reqs):
                                    if r and r[0] == req["REQUEST_ID"]:
                                        aws.update(f"B{i+1}", [[edit_keyed]])
                                        aws.update(f"D{i+1}", [[edit_name]])
                                        aws.update(f"E{i+1}", [[edit_contact]])
                                        aws.update(f"F{i+1}", [[edit_pos]])
                                        aws.update(f"G{i+1}", [[edit_dept]])
                                        aws.update(f"H{i+1}", [[edit_proj]])
                                        aws.update(f"I{i+1}", [[edit_job]])
                                        aws.update(f"J{i+1}", [[edit_loc]])
                                        break
                                st.cache_data.clear()
                                st.success("✅ Request details updated!")
                                st.rerun()

                    except:
                        st.warning("Could not load item details.")
                        items = []

                    approver_name = st.text_input(
                        "Your name (approver) ✱",
                        key=f"approver_{req['REQUEST_ID']}"
                    )
                    ac1, ac2, ac3 = st.columns(3)
                    with ac1:
                        if st.button(f"✅ Approve Request {req['REQUEST_ID']}", type="primary",
                                     key=f"approve_{req['REQUEST_ID']}"):
                            if not approver_name:
                                st.error("❌ Please enter your name to approve.")
                            else:
                                # Find row in approvals sheet
                                all_reqs = aws.get_all_values()
                                for i, r in enumerate(all_reqs):
                                    if r and r[0] == req["REQUEST_ID"]:
                                        aws.update(f"M{i+1}", [["APPROVED"]])
                                        aws.update(f"N{i+1}", [[approver_name]])
                                        break

                                ws = get_worksheet()
                                req_type = req.get("TYPE", "ISSUE") or "ISSUE"

                                if req_type == "RETURN":
                                    # Mark items back as AVAILABLE
                                    for item in items:
                                        idx = item.get("_idx")
                                        if idx is not None:
                                            row_num = int(idx) + 2
                                            try:
                                                ws.update(f"P{row_num}:Y{row_num}", [[
                                                    item.get("DATE OUT",""), req["REQ_DATE"],
                                                    item.get("REQUESTOR",""), item.get("JOB NO",""),
                                                    item.get("PROJECT / USAGE",""), item.get("LOCATION",""),
                                                    item.get("CONDITION OUT",""), "GOOD",
                                                    "AVAILABLE", ""
                                                ]])
                                            except:
                                                pass
                                    st.success(f"✅ Return approved by {approver_name}! Items marked AVAILABLE. Form now available.")
                                else:
                                    # ISSUE — mark items OUT
                                    for item in items:
                                        idx = item.get("_idx")
                                        if idx is not None:
                                            row_num = int(idx) + 2
                                            try:
                                                ws.update(f"P{row_num}:Y{row_num}", [[
                                                    req["REQ_DATE"], "",
                                                    req["REQ_NAME"], req["REQ_JOBNO"],
                                                    req["REQ_PROJECT"], req["REQ_LOCATION"],
                                                    "GOOD", "",
                                                    "OUT", ""
                                                ]])
                                            except:
                                                pass
                                    st.success(f"✅ Approved by {approver_name}! Items marked OUT. PDF now available.")

                                st.cache_data.clear()
                                st.rerun()

                    with ac2:
                        if st.button(f"❌ Reject Request {req['REQUEST_ID']}",
                                     key=f"reject_{req['REQUEST_ID']}"):
                            all_reqs = aws.get_all_values()
                            for i, r in enumerate(all_reqs):
                                if r and r[0] == req["REQUEST_ID"]:
                                    aws.update(f"M{i+1}", [["REJECTED"]])
                                    break
                            st.cache_data.clear()
                            st.warning(f"Request {req['REQUEST_ID']} rejected.")
                            st.rerun()


                    with ac3:
                        del_key = f"confirm_del_{req['REQUEST_ID']}"
                        if del_key not in st.session_state:
                            st.session_state[del_key] = False

                        if not st.session_state[del_key]:
                            if st.button(f"🗑️ Delete", key=f"del_btn_{req['REQUEST_ID']}"):
                                st.session_state[del_key] = True
                                st.rerun()
                        else:
                            st.warning("⚠️ Are you sure? This cannot be undone.")
                            del_by = st.text_input("Your name (deleting this)", key=f"delby_{req['REQUEST_ID']}")
                            reason_options = [
                                "Duplicate entry",
                                "Wrong item selected",
                                "Wrong requestor information",
                                "Request made by mistake",
                                "Test entry / not a real request",
                                "Cancelled by requestor",
                                "Other (type below)"
                            ]
                            del_reason_choice = st.selectbox(
                                "Reason for deletion", reason_options, key=f"delreason_sel_{req['REQUEST_ID']}"
                            )
                            if del_reason_choice == "Other (type below)":
                                del_reason = st.text_input("Type your reason", key=f"delreason_{req['REQUEST_ID']}")
                            else:
                                del_reason = del_reason_choice
                            dc1, dc2 = st.columns(2)
                            with dc1:
                                if st.button("✅ Yes, Delete", key=f"del_confirm_{req['REQUEST_ID']}"):
                                    if not del_by:
                                        st.error("❌ Please enter your name.")
                                    else:
                                        delete_approval_request(req["REQUEST_ID"], del_by, del_reason)
                                        st.session_state[del_key] = False
                                        st.cache_data.clear()
                                        st.success("Request deleted.")
                                        st.rerun()
                            with dc2:
                                if st.button("❌ Cancel", key=f"del_cancel_{req['REQUEST_ID']}"):
                                    st.session_state[del_key] = False
                                    st.rerun()

                 

    # ── APPROVED ──────────────────────────────────────────────────────────────
    with tab_approved:
        approved = adf[adf["STATUS"] == "APPROVED"] if not adf.empty else pd.DataFrame()
        if approved.empty:
            st.info("No approved requests yet.")
        else:
            st.markdown(f"#### {len(approved)} Approved Request(s)")
            for _, req in approved.iterrows():
                with st.expander(
                    f"✅ Request {req['REQUEST_ID']} — {req['REQ_NAME']} | "
                    f"Job: {req['REQ_JOBNO']} | Approved by: {req['APPROVED_BY']}"
                ):
                    st.markdown(f"""
                    **Requestor:** {req['REQ_NAME']} | **Job No:** {req['REQ_JOBNO']}
                    **Project:** {req['REQ_PROJECT']} | **Keyed by:** {req['KEYED_BY']}
                    **Approved by:** {req['APPROVED_BY']} | **Date:** {req['DATE_REQUESTED']}
                    """)

                    try:
                        items = json.loads(req["ITEMS_JSON"])
                        items_df = pd.DataFrame(items)
                        disp = [c for c in ["TAGGING NUMBER","DESCRIPTION","SERIAL NUMBER","BRAND","SIZE"] if c in items_df.columns]
                        st.dataframe(items_df[disp], use_container_width=True)
                        with st.expander(f"✏️ Edit Request Details — {req['REQUEST_ID']}"):
                            ec1, ec2 = st.columns(2)
                            with ec1:
                                edit_name    = st.text_input("Requestor Name", value=req.get("REQ_NAME","") or "", key=f"edit_name_{req['REQUEST_ID']}")
                                edit_contact = st.text_input("Contact No",     value=req.get("REQ_CONTACT","") or "", key=f"edit_contact_{req['REQUEST_ID']}")
                                edit_pos     = st.text_input("Position",      value=req.get("REQ_POSITION","") or "", key=f"edit_pos_{req['REQUEST_ID']}")
                                edit_dept    = st.text_input("Department",    value=req.get("REQ_DEPT","") or "", key=f"edit_dept_{req['REQUEST_ID']}")
                            with ec2:
                                edit_proj    = st.text_input("Project",       value=req.get("REQ_PROJECT","") or "", key=f"edit_proj_{req['REQUEST_ID']}")
                                edit_job     = st.text_input("Job No",        value=req.get("REQ_JOBNO","") or "", key=f"edit_job_{req['REQUEST_ID']}")
                                edit_loc     = st.text_input("Location",      value=req.get("REQ_LOCATION","") or "", key=f"edit_loc_{req['REQUEST_ID']}")
                                edit_keyed   = st.text_input("Keyed By",      value=req.get("KEYED_BY","") or "", key=f"edit_keyed_{req['REQUEST_ID']}")

                            if st.button(f"💾 Save Edits — {req['REQUEST_ID']}", key=f"save_edit_{req['REQUEST_ID']}"):
                                all_reqs = aws.get_all_values()
                                for i, r in enumerate(all_reqs):
                                    if r and r[0] == req["REQUEST_ID"]:
                                        aws.update(f"B{i+1}", [[edit_keyed]])
                                        aws.update(f"D{i+1}", [[edit_name]])
                                        aws.update(f"E{i+1}", [[edit_contact]])
                                        aws.update(f"F{i+1}", [[edit_pos]])
                                        aws.update(f"G{i+1}", [[edit_dept]])
                                        aws.update(f"H{i+1}", [[edit_proj]])
                                        aws.update(f"I{i+1}", [[edit_job]])
                                        aws.update(f"J{i+1}", [[edit_loc]])
                                        break
                                st.cache_data.clear()
                                st.success("✅ Request details updated!")
                                st.rerun()
                    except:
                        items = []

                    # PDF download — only available for approved
                    if st.button(f"📄 Download Issue PDF — {req['REQUEST_ID']}",
                                 key=f"pdf_{req['REQUEST_ID']}"):
                        pdf_data = {
                            "req_name":    req["REQ_NAME"],
                            "req_contact": req["REQ_CONTACT"],
                            "req_position":req["REQ_POSITION"],
                            "req_dept":    req["REQ_DEPT"],
                            "req_project": req["REQ_PROJECT"],
                            "req_jobno":   req["REQ_JOBNO"],
                            "req_location":req["REQ_LOCATION"],
                            "req_date":    req["REQ_DATE"],
                            "req_retdate": req["REQ_RETDATE"],
                            "date_out":    req["REQ_DATE"],
                            "cond_out":    "GOOD",
                            "keyed_by":    req["KEYED_BY"],
                            "approved_by": req["APPROVED_BY"],
                            "items":       items
                        }
                        pdf_bytes = generate_issue_pdf(pdf_data)
                        st.download_button(
                            label=f"⬇️ Download Issue Form — {req['REQUEST_ID']}",
                            data=pdf_bytes,
                            file_name=f"Issue_{req['REQ_NAME']}_{req['REQUEST_ID']}.pdf",
                            mime="application/pdf",
                            key=f"dl_{req['REQUEST_ID']}"
                        )

                    if req.get("TYPE") == "RETURN":
                        first_item = items[0] if items else {}
                        ret_data = {
                            "req_name":     req.get("REQ_NAME",""),
                            "req_jobno":    req.get("REQ_JOBNO",""),
                            "req_project":  req.get("REQ_PROJECT",""),
                            "req_location": req.get("REQ_LOCATION",""),
                            "date_out":     first_item.get("DATE OUT",""),
                            "date_ret":     req.get("REQ_DATE",""),
                            "ret_status":   req.get("RET_STATUS","FULLY RETURNED") or "FULLY RETURNED",
                            "cond_ret":     "GOOD",
                            "remarks":      "",
                            "keyed_by":     req.get("KEYED_BY",""),
                            "approved_by":  req.get("APPROVED_BY",""),
                            "items":        items
                        }
                        if st.button(f"📄 Download Return PDF — {req['REQUEST_ID']}",
                                        key=f"retpdf_{req['REQUEST_ID']}"):
                            pdf_bytes = generate_return_pdf(ret_data)
                            st.download_button(
                                label=f"⬇️ Download Return Form PDF — {req['REQUEST_ID']}",
                                data=pdf_bytes,
                                file_name=f"Return_{req['REQ_NAME']}_{req['REQUEST_ID']}.pdf",
                                mime="application/pdf",
                                key=f"retdl_{req['REQUEST_ID']}"
                            )

                        if st.button(f"📝 Download Return Word — {req['REQUEST_ID']}",
                                        key=f"retdocx_{req['REQUEST_ID']}"):
                            docx_bytes = generate_return_docx(ret_data)
                            st.download_button(
                                label=f"⬇️ Download Return Form Word — {req['REQUEST_ID']}",
                                data=docx_bytes,
                                file_name=f"Return_{req['REQ_NAME']}_{req['REQUEST_ID']}.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                key=f"retdocxdl_{req['REQUEST_ID']}"
                            )

                    if st.button(f"📝 Download Issue Word — {req['REQUEST_ID']}",
                                key=f"docx_{req['REQUEST_ID']}"):
                        docx_data = {
                            "req_name":    req["REQ_NAME"],
                            "req_contact": req["REQ_CONTACT"],
                            "req_position":req["REQ_POSITION"],
                            "req_dept":    req["REQ_DEPT"],
                            "req_project": req["REQ_PROJECT"],
                            "req_jobno":   req["REQ_JOBNO"],
                            "req_location":req["REQ_LOCATION"],
                            "req_date":    req["REQ_DATE"],
                            "req_retdate": req["REQ_RETDATE"],
                            "date_out":    req["REQ_DATE"],
                            "cond_out":    "GOOD",
                            "keyed_by":    req["KEYED_BY"],
                            "approved_by": req["APPROVED_BY"],
                            "items":       items
                        }
                        docx_bytes = generate_issue_docx(docx_data)
                        st.download_button(
                            label=f"⬇️ Download Issue Word — {req['REQUEST_ID']}",
                            data=docx_bytes,
                            file_name=f"Issue_{req['REQ_NAME']}_{req['REQUEST_ID']}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            key=f"docxdl_{req['REQUEST_ID']}"
                        )   
                    st.markdown("---")
                    del_key2 = f"confirm_del2_{req['REQUEST_ID']}"
                    if del_key2 not in st.session_state:
                        st.session_state[del_key2] = False

                    if not st.session_state[del_key2]:
                        if st.button(f"🗑️ Delete This Request", key=f"del_btn2_{req['REQUEST_ID']}"):
                            st.session_state[del_key2] = True
                            st.rerun()
                    else:
                        st.warning("⚠️ Are you sure? This cannot be undone.")
                        del_by2 = st.text_input("Your name (deleting this)", key=f"delby2_{req['REQUEST_ID']}")
                        reason_options2 = [
                            "Duplicate entry",
                            "Wrong item selected",
                            "Wrong requestor information",
                            "Request made by mistake",
                            "Test entry / not a real request",
                            "Cancelled by requestor",
                            "Other (type below)"
                        ]
                        del_reason_choice2 = st.selectbox(
                            "Reason for deletion", reason_options2, key=f"delreason_sel2_{req['REQUEST_ID']}"
                        )
                        if del_reason_choice2 == "Other (type below)":
                            del_reason2 = st.text_input("Type your reason", key=f"delreason2_{req['REQUEST_ID']}")
                        else:
                            del_reason2 = del_reason_choice2
                        dc1, dc2 = st.columns(2)
                        with dc1:
                            if st.button("✅ Yes, Delete", key=f"del_confirm2_{req['REQUEST_ID']}"):
                                if not del_by2:
                                    st.error("❌ Please enter your name.")
                                else:
                                    delete_approval_request(req["REQUEST_ID"], del_by2, del_reason2)
                                    st.session_state[del_key2] = False
                                    st.cache_data.clear()
                                    st.success("Request deleted.")
                                    st.rerun()
                        with dc2:
                            if st.button("❌ Cancel", key=f"del_cancel2_{req['REQUEST_ID']}"):
                                st.session_state[del_key2] = False
                                st.rerun()

# ══════════════════════════════════════════════════════════════════════════════
# PAGE 6 — SETTINGS
# ══════════════════════════════════════════════════════════════════════════════
elif page == "⚙️ Settings":
    st.markdown("## ⚙️ Settings")
    st.markdown("### Manage Category Dropdown List")
    st.caption("Add or remove categories available in the Add/Edit Item forms.")

    categories = load_categories()

    st.markdown("---")
    st.markdown(f"**Current Categories ({len(categories)})**")

    if not categories:
        st.info("No categories yet. Add one below.")
    else:
        for cat in categories:
            c1, c2 = st.columns([5, 1])
            with c1:
                st.write(cat)
            with c2:
                del_key = f"confirm_del_cat_{cat}"
                if del_key not in st.session_state:
                    st.session_state[del_key] = False

                if not st.session_state[del_key]:
                    if st.button("✕", key=f"del_cat_{cat}"):
                        st.session_state[del_key] = True
                        st.rerun()

        # Show confirmation for whichever category was clicked
        for cat in categories:
            del_key = f"confirm_del_cat_{cat}"
            if st.session_state.get(del_key):
                st.warning(f"⚠️ Are you sure you want to remove **{cat}**? This cannot be undone.")
                cc1, cc2 = st.columns(2)
                with cc1:
                    if st.button(f"✅ Yes, remove {cat}", key=f"confirm_yes_{cat}"):
                        remove_category(cat)
                        st.session_state[del_key] = False
                        st.success(f"Removed '{cat}'")
                        st.rerun()
                with cc2:
                    if st.button(f"❌ Cancel", key=f"confirm_no_{cat}"):
                        st.session_state[del_key] = False
                        st.rerun()

    st.markdown("---")
    st.markdown("**Add New Category**")
    new_cat_input = st.text_input("Type new category name", key="new_cat_input")
    if st.button("➕ Add Category", type="primary"):
        if not new_cat_input.strip():
            st.error("❌ Please type a category name.")
        else:
            success = add_category(new_cat_input)
            if success:
                st.success(f"✅ Added '{new_cat_input.strip().upper()}'")
                st.rerun()
            else:
                st.error("❌ This category already exists.")